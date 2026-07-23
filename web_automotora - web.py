import streamlit as st
import psycopg2
from psycopg2 import errors as pg_errors
from docx import Document
import os
import io
import hashlib
import secrets
import pandas as pd

# ----------------------------------------------------------------
# CONFIGURACIÓN DE LA PÁGINA
# ----------------------------------------------------------------
st.set_page_config(
    page_title="San Martín Automotriz",
    page_icon="🚗",
    layout="wide"
)

# ----------------------------------------------------------------
# VALIDADOR Y FORMATEADOR DE RUT CHILENO
# ----------------------------------------------------------------
def validar_rut(rut):
    if not rut or rut.strip() == "":
        return True, ""
        
    rut_limpio = str(rut).replace(".", "").replace("-", "").replace(" ", "").upper()
    
    if len(rut_limpio) < 8 or len(rut_limpio) > 9:
        return False, rut
        
    cuerpo = rut_limpio[:-1]
    dv_ingresado = rut_limpio[-1]
    
    if not cuerpo.isdigit():
        return False, rut
        
    suma = 0
    multiplo = 2
    for d in reversed(cuerpo):
        suma += int(d) * multiplo
        multiplo += 1
        if multiplo == 8:
            multiplo = 2
            
    resto = suma % 11
    dv_calculado = 11 - resto
    
    if dv_calculado == 11:
        dv_esperado = "0"
    elif dv_calculado == 10:
        dv_esperado = "K"
    else:
        dv_esperado = str(dv_calculado)
        
    if dv_ingresado == dv_esperado:
        return True, f"{cuerpo}-{dv_esperado}"
    else:
        return False, rut

# ----------------------------------------------------------------
# REEMPLAZO DE MARCADORES [CLAVE] EN PLANTILLAS WORD
# ----------------------------------------------------------------
def reemplazar_en_parrafo(parrafo, mapa):
    """Reemplaza cada [CLAVE] por su valor, incluso si Word partió el marcador
    en varios "runs" internos (pasa seguido al editar una plantilla). Para
    lograrlo, se recalcula el texto completo del párrafo y se vuelve a
    escribir en el primer run, vaciando el resto — por eso el párrafo puede
    perder variaciones de formato (negrita/cursiva) que tuviera a mitad de
    frase, aunque conserva el estilo general del primer run."""
    texto_original = parrafo.text
    texto_nuevo = texto_original
    for clave, valor in mapa.items():
        if clave in texto_nuevo:
            texto_nuevo = texto_nuevo.replace(clave, str(valor) if valor else "")
    if texto_nuevo == texto_original or not parrafo.runs:
        return
    parrafo.runs[0].text = texto_nuevo
    for run in parrafo.runs[1:]:
        run.text = ""

# ----------------------------------------------------------------
# CONEXIÓN A LA BASE DE DATOS
# ----------------------------------------------------------------
def conectar_bd():
    return psycopg2.connect(
        dbname=st.secrets["DB_NAME"],
        user=st.secrets["DB_USER"],
        password=st.secrets["DB_PASS"],
        host=st.secrets["DB_HOST"],
        port="5432"
    )

# ----------------------------------------------------------------
# FORMATEO AUTOMÁTICO DE MONTOS (agrega puntos de miles al salir del campo)
# ----------------------------------------------------------------
def formatear_monto_callback(key):
    digitos = "".join(c for c in st.session_state.get(key, "") if c.isdigit())
    st.session_state[key] = f"{int(digitos):,}".replace(",", ".") if digitos else "0"

# ----------------------------------------------------------------
# USUARIOS Y PERMISOS
# ----------------------------------------------------------------
TABS_DISPONIBLES = [
    ("documentos", "📝 Generar Documentos"),
    ("productos", "📦 Productos"),
    ("clientes", "👤 Clientes"),
    ("creditos", "💳 Créditos"),
    ("vehiculos", "🚘 Vehículos"),
    ("buscar", "🔍 Buscar"),
]

def hash_password(password, salt=None):
    if salt is None:
        salt = secrets.token_hex(16)
    hash_resultante = hashlib.sha256((salt + password).encode()).hexdigest()
    return hash_resultante, salt

def verificar_password(password, salt, hash_guardado):
    hash_calculado, _ = hash_password(password, salt)
    return hash_calculado == hash_guardado

def asegurar_tabla_usuarios():
    """Crea la tabla de usuarios si no existe y migra el admin inicial la primera vez."""
    con = conectar_bd()
    cur = con.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS usuarios (
            id                   SERIAL PRIMARY KEY,
            username             VARCHAR(50)  NOT NULL UNIQUE,
            password_hash        VARCHAR(128) NOT NULL,
            password_salt        VARCHAR(64)  NOT NULL,
            nombre_completo      VARCHAR(255),
            perfil               VARCHAR(50)  NOT NULL DEFAULT 'Usuario',
            es_admin             BOOLEAN      NOT NULL DEFAULT FALSE,
            pestañas_permitidas  TEXT,
            activo               BOOLEAN      NOT NULL DEFAULT TRUE,
            creado_en            TIMESTAMP    NOT NULL DEFAULT now()
        )
    """)
    cur.execute("SELECT COUNT(*) FROM usuarios")
    if cur.fetchone()[0] == 0:
        hash_admin, salt_admin = hash_password("automotora2026")
        todas = ",".join(k for k, _ in TABS_DISPONIBLES)
        cur.execute("""
            INSERT INTO usuarios (username, password_hash, password_salt, nombre_completo, perfil, es_admin, pestañas_permitidas)
            VALUES (%s,%s,%s,%s,%s,%s,%s)
        """, ("sanmartin", hash_admin, salt_admin, "Administrador", "Administrador", True, todas))
    con.commit()
    cur.close()
    con.close()

# ================================================================
# SISTEMA DE LOGIN (SEGURIDAD)
# ================================================================
if "autenticado" not in st.session_state:
    st.session_state["autenticado"] = False

if not st.session_state["autenticado"]:
    asegurar_tabla_usuarios()
    st.markdown("<br><br><h2 style='text-align: center;'>🔒 Acceso Restringido</h2>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: gray;'>Por favor, ingresa tus credenciales para continuar.</p>", unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        with st.form("form_login"):
            usuario = st.text_input("Nombre de Usuario")
            clave = st.text_input("Contraseña", type="password")
            ingresar = st.form_submit_button("🚪 Iniciar Sesión", use_container_width=True)

            if ingresar:
                try:
                    con = conectar_bd()
                    cur = con.cursor()
                    cur.execute(
                        "SELECT id, password_hash, password_salt, nombre_completo, perfil, es_admin, pestañas_permitidas, activo "
                        "FROM usuarios WHERE username=%s",
                        (usuario.strip(),),
                    )
                    fila = cur.fetchone()
                    cur.close()
                    con.close()
                except Exception as e:
                    fila = None
                    st.error(f"❌ Error de conexión: {e}")

                if fila and fila[7] and verificar_password(clave, fila[2], fila[1]):
                    st.session_state["autenticado"] = True
                    st.session_state["usuario_actual"] = {
                        "id": fila[0],
                        "username": usuario.strip(),
                        "nombre_completo": fila[3],
                        "perfil": fila[4],
                        "es_admin": fila[5],
                        "pestañas_permitidas": set((fila[6] or "").split(",")),
                    }
                    st.rerun()
                elif fila and not fila[7]:
                    st.error("❌ Este usuario está desactivado.")
                else:
                    st.error("❌ Usuario o contraseña incorrectos.")
    st.stop()

# ================================================================
# DISEÑO VISUAL Y TÍTULO
# ================================================================
st.markdown("""
    <style>
    .titulo-principal {
        font-size: 3rem !important;
        font-weight: 800 !important;
        color: #1E3A8A;
        text-align: center;
        background: linear-gradient(90deg, #1E3A8A, #3B82F6);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 0px;
        padding-bottom: 10px;
    }
    .subtitulo {
        text-align: center;
        color: #64748B;
        font-size: 1.2rem;
        margin-bottom: 30px;
    }
    </style>
""", unsafe_allow_html=True)

st.markdown('<h1 class="titulo-principal">Sistema San Martín Automotriz</h1>', unsafe_allow_html=True)
st.markdown('<p class="subtitulo">Plataforma Integral de Gestión y Documentos</p>', unsafe_allow_html=True)

# Mensajes de éxito pendientes de mostrar tras un st.rerun() (evita que el
# mensaje se pierda porque el rerun ocurre antes de que el navegador lo pinte).
if st.session_state.get("_flash_ok"):
    st.success(st.session_state.pop("_flash_ok"))

# ----------------------------------------------------------------
# BARRA DE SESIÓN
# ----------------------------------------------------------------
usuario_actual = st.session_state["usuario_actual"]
sc1, sc2 = st.columns([5, 1])
with sc1:
    st.caption(f"👤 Conectado como **{usuario_actual['username']}** ({usuario_actual['perfil']})")
with sc2:
    if st.button("🚪 Cerrar Sesión", use_container_width=True):
        for k in list(st.session_state.keys()):
            del st.session_state[k]
        st.rerun()

# ----------------------------------------------------------------
# CREACIÓN DE PESTAÑAS PRINCIPALES (según permisos del usuario)
# ----------------------------------------------------------------
es_admin = usuario_actual["es_admin"]
permitidas = usuario_actual["pestañas_permitidas"]

tabs_visibles = [(k, label) for k, label in TABS_DISPONIBLES if es_admin or k in permitidas]
if es_admin:
    tabs_visibles.append(("admin", "⚙️ Administración"))

_tab_objs = st.tabs([label for _, label in tabs_visibles])
tabs = dict(zip([k for k, _ in tabs_visibles], _tab_objs))

FORMAS_PAGO_DISPONIBLES = ["Efectivo", "Tarjeta Redcompra", "Tarjeta de Crédito", "Transferencia"]
IVA_RATE = 0.19

# ================================================================
# PESTAÑA 1 — GENERAR DOCUMENTOS
# ================================================================
@st.fragment
def render_documentos():
    st.markdown("### 📄 Selecciona el Documento a Generar")
    
    sub_compraventa, sub_consignacion, sub_declaracion, sub_poder, sub_notarial = st.tabs([
        "🤝 Compraventa", 
        "📋 Consignación", 
        "⚖️ Declaración Jurada", 
        "📜 Carta Poder", 
        "🏛️ Contrato Notarial"
    ])

    # ---------------------------------------------------------
    # MÓDULO 1: COMPRAVENTA
    # ---------------------------------------------------------
    with sub_compraventa:
        if "menu_compraventa" not in st.session_state:
            st.session_state["menu_compraventa"] = "crear"

        st.markdown("**¿Qué deseas hacer con la Compraventa?**")
        col_btn1, col_btn2 = st.columns(2)
        with col_btn1:
            if st.button("➕ Crear y guardar nuevo contrato", key="btn_c_crear", use_container_width=True):
                st.session_state["menu_compraventa"] = "crear"
        with col_btn2:
            if st.button("✏️ Modificar / Reimprimir existente", key="btn_c_mod", use_container_width=True):
                st.session_state["menu_compraventa"] = "modificar"
        st.write("") 

        if st.session_state["menu_compraventa"] == "crear":
            st.subheader("Crear Nuevo Contrato de Compraventa")
            with st.form("form_nuevo_contrato", clear_on_submit=False):
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("📝 **Datos de la Venta y Comprador**")
                    numero_contrato_val = st.text_input("Número de Contrato *", placeholder="Ej: 1539")
                    fecha_val = st.text_input("Fecha del Contrato", placeholder="Ej: 19.05.2026")
                    nombre_val = st.text_input("Nombre del Comprador *", placeholder="Ej: EDUARDO PALMA")
                    rut_val = st.text_input("RUT del Comprador", placeholder="Ej: 16.328.681-5")
                    nacionalidad_val = st.text_input("Nacionalidad", placeholder="Ej: chilena")
                    direccion_val = st.text_input("Dirección", placeholder="Ej: SECTOR HUILLINCO S/N")
                    comuna_val = st.text_input("Comuna/Ciudad", placeholder="Ej: COMUNA CHONCHI, CHILOE")
                    
                    st.markdown("👤 **Datos del Vendedor / Automotora**")
                    nombre_vendedor_val = st.text_input("Nombre del Vendedor", placeholder="Ej: JORGE MALDONADO LOBO")
                    rut_vendedor_val = st.text_input("RUT del Vendedor", placeholder="Ej: 7.790.477-8")

                    st.markdown("💰 **Datos Financieros**")
                    precio_num_val = st.text_input("Precio en Números", placeholder="Ej: $10.990.000.-")
                    precio_pal_val = st.text_input("Precio en Palabras", placeholder="Ej: DIEZ MILLONES NOVECIENTOS NOVENTA MIL")
                    forma_pago_val = st.text_input("Forma de Pago", placeholder="Ej: EFECTIVO / TRANSFERENCIA")
                    gastos_val = st.text_input("Gastos Notariales", value="$294.850")

                with col2:
                    st.markdown("🚗 **Datos del Vehículo**")
                    tipo_veh_val = st.text_input("Tipo de Vehículo", placeholder="Ej: CAMIONETA")
                    marca_val = st.text_input("Marca", placeholder="Ej: NISSAN")
                    modelo_val = st.text_input("Modelo", placeholder="Ej: TERRANO 2.5 4X4")
                    anio_val = st.text_input("Año", placeholder="Ej: 2011")
                    motor_val = st.text_input("N° Motor", placeholder="Ej: YD25257334T")
                    chasis_val = st.text_input("N° Chasis", placeholder="Ej: 3N63N6PD21Y2ZK878230")
                    n_vin_val = st.text_input("N° VIN", placeholder="Ej: 3N63N6PD21Y2ZK878230")
                    color_val = st.text_input("Color", placeholder="Ej: ROJO")
                    patente_val = st.text_input("Patente", placeholder="Ej: CXLJ.20-8")
                    
                    st.markdown("ℹ️ **Información Adicional**")
                    observaciones_val = st.text_area("Observaciones (Obs)", placeholder="Ej: Detalles estéticos menores...")

                generar = st.form_submit_button("💾 Guardar y Generar Contrato")

            if generar:
                if not numero_contrato_val or not nombre_val:
                    st.error("⚠️ El Número de Contrato y el Nombre son obligatorios.")
                else:
                    v_comp, r_comp = validar_rut(rut_val)
                    v_vend, r_vend = validar_rut(rut_vendedor_val)
                    
                    if not v_comp: st.error("❌ El RUT del Comprador es inválido. Por favor, corrígelo.")
                    elif not v_vend: st.error("❌ El RUT del Vendedor es inválido. Por favor, corrígelo.")
                    else:
                        try:
                            con = conectar_bd()
                            cur = con.cursor()
                            cur.execute("""
                                INSERT INTO contratos (
                                    numero_contrato, fecha_contrato, nombre_comprador, rut_comprador, nacionalidad, 
                                    direccion, comuna, precio_numeros, precio_palabras, gastos_notariales, 
                                    marca, modelo, anio, motor, chasis, color, patente,
                                    nombre_vendedor, rut_vendedor, tipo_vehiculo, n_vin, observaciones, forma_pago
                                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                            """, (
                                numero_contrato_val, fecha_val, nombre_val, r_comp, nacionalidad_val,
                                direccion_val, comuna_val, precio_num_val, precio_pal_val, gastos_val,
                                marca_val, modelo_val, anio_val, motor_val, chasis_val, color_val, patente_val,
                                nombre_vendedor_val, r_vend, tipo_veh_val, n_vin_val, observaciones_val, forma_pago_val
                            ))
                            con.commit()

                            ruta_plantilla = "Plantillas/Plantilla_Compraventa.docx"
                            if not os.path.exists(ruta_plantilla):
                                st.error(f"❌ Falta el archivo {ruta_plantilla}.")
                            else:
                                doc = Document(ruta_plantilla)
                                reemplazos = {
                                    "[NÚMERO_CONTRATO]": numero_contrato_val, "[FECHA_CONTRATO]": fecha_val,
                                    "[NOMBRE_COMPRADOR]": nombre_val, "[RUT_COMPRADOR]": r_comp,
                                    "[NACIONALIDAD]": nacionalidad_val, "[DIRECCION]": direccion_val,
                                    "[COMUNA]": comuna_val, "[PRECIO_NUMEROS]": precio_num_val,
                                    "[PRECIO_PALABRAS]": precio_pal_val, "[GASTOS_NOTARIALES]": gastos_val,
                                    "[MARCA]": marca_val, "[MODELO]": modelo_val, "[AÑO]": anio_val,
                                    "[MOTOR]": motor_val, "[CHASIS]": chasis_val, "[COLOR]": color_val, "[PATENTE]": patente_val,
                                    "[NOMBRE_VENDEDOR]": nombre_vendedor_val, "[RUT_VENDEDOR]": r_vend,
                                    "[TIPO_VEHICULO]": tipo_veh_val, "[N_VIN]": n_vin_val, 
                                    "[FORMA_PAGO]": forma_pago_val, "[OBSERVACIONES]": observaciones_val
                                }
                                for p in doc.paragraphs: reemplazar_en_parrafo(p, reemplazos)
                                for tabla in doc.tables:
                                    for fila in tabla.rows:
                                        for celda in fila.cells:
                                            for p in celda.paragraphs: reemplazar_en_parrafo(p, reemplazos)
                                buffer = io.BytesIO()
                                doc.save(buffer)
                                buffer.seek(0)
                                st.success(f"✅ ¡Contrato N°{numero_contrato_val} guardado y generado con éxito!")
                                st.download_button("📥 Descargar Documento", data=buffer, 
                                                file_name=f"Contrato_{numero_contrato_val}_{nombre_val}.docx", 
                                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        except pg_errors.UniqueViolation:
                            st.error("❌ Ese Número de Contrato ya existe. Ve a la opción 'Modificar'.")
                        except Exception as e:
                            st.error(f"❌ Error: {e}")
                        finally:
                            if "cur" in locals(): cur.close()
                            if "con" in locals(): con.close()

        elif st.session_state["menu_compraventa"] == "modificar":
            st.subheader("Modificar Contrato Existente")
            try:
                con = conectar_bd()
                cur = con.cursor()
                cur.execute("SELECT numero_contrato, nombre_comprador, rut_comprador, patente FROM contratos ORDER BY numero_contrato DESC LIMIT 10")
                filas = cur.fetchall()
                if filas:
                    st.markdown("📋 **Últimos Contratos guardados:**")
                    st.dataframe(pd.DataFrame(filas, columns=["N° Doc", "Nombre", "RUT", "Patente"]), use_container_width=True, hide_index=True)
            except: pass
            finally:
                if "cur" in locals(): cur.close()
                if "con" in locals(): con.close()

            st.write("---")
            nro_mod = st.text_input("Ingresa el N° de Contrato a buscar:", key="nro_modificar_c")
            if st.button("🔎 Buscar", key="btn_busc_compra"):
                try:
                    con = conectar_bd()
                    cur = con.cursor()
                    cur.execute("SELECT * FROM contratos WHERE numero_contrato = %s", (nro_mod.strip(),))
                    cols = [d[0] for d in cur.description]
                    fila = cur.fetchone()
                    if fila:
                        st.session_state["contrato_mod"] = dict(zip(cols, fila))
                        st.success("✅ Contrato encontrado. Edita los campos abajo y vuelve a generar.")
                    else: st.error("❌ No se encontró.")
                except Exception as e: st.error(f"Error: {e}")
                finally:
                    if "cur" in locals(): cur.close()
                    if "con" in locals(): con.close()

            if "contrato_mod" in st.session_state:
                c = st.session_state["contrato_mod"]
                with st.form("form_mod_contrato"):
                    col1, col2 = st.columns(2)
                    with col1:
                        fecha_m = st.text_input("Fecha", value=c.get("fecha_contrato",""))
                        nombre_m = st.text_input("Nombre Comprador", value=c.get("nombre_comprador",""))
                        rut_m = st.text_input("RUT Comprador", value=c.get("rut_comprador",""))
                        nac_m = st.text_input("Nacionalidad", value=c.get("nacionalidad",""))
                        dir_m = st.text_input("Dirección", value=c.get("direccion",""))
                        com_m = st.text_input("Comuna", value=c.get("comuna",""))
                        vendedor_nom_m = st.text_input("Nombre Vendedor", value=c.get("nombre_vendedor",""))
                        vendedor_rut_m = st.text_input("RUT Vendedor", value=c.get("rut_vendedor",""))
                        p_num_m = st.text_input("Precio Números", value=c.get("precio_numeros",""))
                        p_pal_m = st.text_input("Precio Palabras", value=c.get("precio_palabras",""))
                        f_pago_m = st.text_input("Forma de Pago", value=c.get("forma_pago",""))
                        gas_m = st.text_input("Gastos Notariales", value=c.get("gastos_notariales",""))
                    with col2:
                        tipo_veh_m = st.text_input("Tipo Vehículo", value=c.get("tipo_vehiculo",""))
                        marca_m = st.text_input("Marca", value=c.get("marca",""))
                        modelo_m = st.text_input("Modelo", value=c.get("modelo",""))
                        anio_m = st.text_input("Año", value=c.get("anio",""))
                        motor_m = st.text_input("N° Motor", value=c.get("motor",""))
                        chasis_m = st.text_input("N° Chasis", value=c.get("chasis",""))
                        n_vin_m = st.text_input("N° VIN", value=c.get("n_vin",""))
                        color_m = st.text_input("Color", value=c.get("color",""))
                        pat_m = st.text_input("Patente", value=c.get("patente",""))
                        obs_m = st.text_area("Observaciones", value=c.get("observaciones",""))
                    actualizar_y_generar = st.form_submit_button("💾 Actualizar y Reimprimir Word")

                if actualizar_y_generar:
                    v_comp, r_comp = validar_rut(rut_m)
                    v_vend, r_vend = validar_rut(vendedor_rut_m)
                    
                    if not v_comp: st.error("❌ RUT Comprador inválido.")
                    elif not v_vend: st.error("❌ RUT Vendedor inválido.")
                    else:
                        try:
                            con = conectar_bd()
                            cur = con.cursor()
                            cur.execute("""
                                UPDATE contratos SET
                                    fecha_contrato=%s, nombre_comprador=%s, rut_comprador=%s, nacionalidad=%s, 
                                    direccion=%s, comuna=%s, precio_numeros=%s, precio_palabras=%s, gastos_notariales=%s, 
                                    marca=%s, modelo=%s, anio=%s, motor=%s, chasis=%s, color=%s, patente=%s,
                                    nombre_vendedor=%s, rut_vendedor=%s, tipo_vehiculo=%s, n_vin=%s, observaciones=%s, forma_pago=%s
                                WHERE numero_contrato=%s
                            """, (
                                fecha_m, nombre_m, r_comp, nac_m, dir_m, com_m, p_num_m, p_pal_m, gas_m,
                                marca_m, modelo_m, anio_m, motor_m, chasis_m, color_m, pat_m,
                                vendedor_nom_m, r_vend, tipo_veh_m, n_vin_m, obs_m, f_pago_m,
                                c["numero_contrato"]
                            ))
                            con.commit()
                            ruta_plantilla = "Plantillas/Plantilla_Compraventa.docx"
                            doc = Document(ruta_plantilla)
                            reemplazos = {
                                "[NÚMERO_CONTRATO]": c["numero_contrato"], "[FECHA_CONTRATO]": fecha_m,
                                "[NOMBRE_COMPRADOR]": nombre_m, "[RUT_COMPRADOR]": r_comp,
                                "[NACIONALIDAD]": nac_m, "[DIRECCION]": dir_m, "[COMUNA]": com_m, 
                                "[PRECIO_NUMEROS]": p_num_m, "[PRECIO_PALABRAS]": p_pal_m, "[GASTOS_NOTARIALES]": gas_m,
                                "[MARCA]": marca_m, "[MODELO]": modelo_m, "[AÑO]": anio_m, "[MOTOR]": motor_m, "[CHASIS]": chasis_m, 
                                "[COLOR]": color_m, "[PATENTE]": pat_m, "[NOMBRE_VENDEDOR]": vendedor_nom_m, "[RUT_VENDEDOR]": r_vend,
                                "[TIPO_VEHICULO]": tipo_veh_m, "[N_VIN]": n_vin_m, "[FORMA_PAGO]": f_pago_m, "[OBSERVACIONES]": obs_m
                            }
                            for p in doc.paragraphs: reemplazar_en_parrafo(p, reemplazos)
                            for tabla in doc.tables:
                                for fila in tabla.rows:
                                    for celda in fila.cells:
                                        for p in celda.paragraphs: reemplazar_en_parrafo(p, reemplazos)
                            buffer = io.BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.success("✅ ¡Contrato actualizado y reimpreso!")
                            st.download_button("📥 Descargar Documento", data=buffer, 
                                            file_name=f"Contrato_{c['numero_contrato']}_Modificado.docx", 
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                            del st.session_state["contrato_mod"]
                        except Exception as e: st.error(f"❌ Error: {e}")
                        finally:
                            if "cur" in locals(): cur.close()
                            if "con" in locals(): con.close()

    # ---------------------------------------------------------
    # MÓDULO 2: CONSIGNACIONES
    # ---------------------------------------------------------
    with sub_consignacion:
        if "menu_consignacion" not in st.session_state:
            st.session_state["menu_consignacion"] = "crear"

        st.markdown("**¿Qué deseas hacer con la Consignación?**")
        col_btn3, col_btn4 = st.columns(2)
        with col_btn3:
            if st.button("➕ Crear y guardar", key="btn_con_crear", use_container_width=True):
                st.session_state["menu_consignacion"] = "crear"
        with col_btn4:
            if st.button("✏️ Modificar / Reimprimir", key="btn_con_mod", use_container_width=True):
                st.session_state["menu_consignacion"] = "modificar"
        st.write("")

        if st.session_state["menu_consignacion"] == "crear":
            st.subheader("Crear Nuevo Contrato de Consignación")
            with st.form("form_nueva_consignacion", clear_on_submit=False):
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("📝 **Datos del Consignante**")
                    nro_consig_val = st.text_input("Número de Consignación *", placeholder="Ej: 001")
                    fecha_consig_val = st.text_input("Fecha", placeholder="Ej: 30 de Mayo del 2026")
                    nombre_consig_val = st.text_input("Nombre Consignante *", placeholder="Ej: SERVICIOS AQUASOL")
                    rut_consig_val = st.text_input("RUT Consignante", placeholder="Ej: 76.053.228-2")
                    dir_consig_val = st.text_input("Dirección")
                    comuna_consig_val = st.text_input("Comuna")
                    tel_consig_val = st.text_input("Teléfono")
                    correo_consig_val = st.text_input("Correo Electrónico")
                    
                    st.markdown("💰 **Datos Financieros y Contrato**")
                    precio_consig_val = st.text_input("Precio de Venta")
                    m_escrito_val = st.text_input("Precio en Palabras")
                    comision_consig_val = st.text_input("Comisión %")
                    d_contrato_val = st.text_input("Duración del Contrato")

                    st.markdown("🏦 **Datos Cuenta Bancaria Consignante**")
                    pago_consignante_val = st.text_input("Método de Pago")
                    cb_nombre_val = st.text_input("Nombre Titular Cuenta")
                    cb_rut_val = st.text_input("RUT Cuenta")
                    cb_banco_val = st.text_input("Banco")
                    cb_cuenta_val = st.text_input("Tipo Cuenta")
                    cb_numero_val = st.text_input("Número de Cuenta")

                with col2:
                    st.markdown("🚗 **Datos del Vehículo**")
                    tipo_veh_val = st.text_input("Tipo de Vehículo")
                    marca_c_val = st.text_input("Marca")
                    modelo_c_val = st.text_input("Modelo")
                    anio_c_val = st.text_input("Año")
                    motor_c_val = st.text_input("N° Motor")
                    chasis_c_val = st.text_input("N° Chasis / VIN")
                    color_c_val = st.text_input("Color")
                    patente_c_val = st.text_input("Patente")
                    kms_c_val = st.text_input("Kilometraje")
                    
                    st.markdown("ℹ️ **Información Adicional**")
                    observaciones_val = st.text_area("Observaciones")

                generar_consig = st.form_submit_button("💾 Guardar y Generar Consignación")

            if generar_consig:
                if not nro_consig_val or not nombre_consig_val:
                    st.error("⚠️ El Número y el Nombre son obligatorios.")
                else:
                    v_consig, r_consig = validar_rut(rut_consig_val)

                    if not v_consig: st.error("❌ RUT del Consignante inválido.")
                    else:
                        try:
                            con = conectar_bd()
                            cur = con.cursor()
                            cur.execute("""
                                INSERT INTO consignaciones (
                                    numero_consignacion, fecha, nombre_consignante, rut_consignante, 
                                    direccion, comuna, telefono, correo, precio, m_escrito, comision, 
                                    d_contrato, pago_consignante, cb_nombre, cb_rut, cb_banco, 
                                    cb_cuenta, cb_numero, tipo_vehiculo, marca, modelo, anio, 
                                    motor, chasis, color, patente, kms, observaciones
                                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                            """, (
                                nro_consig_val, fecha_consig_val, nombre_consig_val, r_consig, 
                                dir_consig_val, comuna_consig_val, tel_consig_val, correo_consig_val, precio_consig_val, m_escrito_val, comision_consig_val,
                                d_contrato_val, pago_consignante_val, cb_nombre_val, cb_rut_val, cb_banco_val,
                                cb_cuenta_val, cb_numero_val, tipo_veh_val, marca_c_val, modelo_c_val, anio_c_val, 
                                motor_c_val, chasis_c_val, color_c_val, patente_c_val, kms_c_val, observaciones_val
                            ))
                            con.commit()

                            ruta_plantilla = "Plantillas/Contrato_consignacion.docx"
                            doc = Document(ruta_plantilla)
                            reemplazos = {
                                "[NÚMERO_CONSIGNACION]": nro_consig_val, "[FECHA]": fecha_consig_val,
                                "[NOMBRE_CONSIGNANTE]": nombre_consig_val, "[RUT_CONSIGNANTE]": r_consig,
                                "[DIRECCION]": dir_consig_val, "[COMUNA]": comuna_consig_val, "[TELEFONO]": tel_consig_val, 
                                "[CORREO]": correo_consig_val, "[PRECIO]": precio_consig_val, "[M_ESCRITO]": m_escrito_val,
                                "[COMISION]": comision_consig_val, "[D_CONTRATO]": d_contrato_val, "[PAGO_CONSIGNANTE]": pago_consignante_val,
                                "[CB_NOMBRE]": cb_nombre_val, "[CB_RUT]": cb_rut_val, "[CB_BANCO]": cb_banco_val,
                                "[CB_CUENTA]": cb_cuenta_val, "[CB_NUMERO]": cb_numero_val, "[TIPO_VEHICULO]": tipo_veh_val,
                                "[MARCA]": marca_c_val, "[MODELO]": modelo_c_val, "[AÑO]": anio_c_val, 
                                "[MOTOR]": motor_c_val, "[CHASIS]": chasis_c_val, "[COLOR]": color_c_val, 
                                "[PATENTE]": patente_c_val, "[KMS]": kms_c_val, "[OBSERVACIONES]": observaciones_val
                            }
                            for p in doc.paragraphs: reemplazar_en_parrafo(p, reemplazos)
                            for tabla in doc.tables:
                                for fila in tabla.rows:
                                    for celda in fila.cells:
                                        for p in celda.paragraphs: reemplazar_en_parrafo(p, reemplazos)
                            buffer = io.BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.success(f"✅ ¡Consignación guardada y generada!")
                            st.download_button("📥 Descargar", data=buffer, file_name=f"Consignacion_{nro_consig_val}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        except pg_errors.UniqueViolation: st.error("❌ Ese Número ya existe.")
                        except Exception as e: st.error(f"❌ Error: {e}")
                        finally:
                            if "cur" in locals(): cur.close()
                            if "con" in locals(): con.close()

        elif st.session_state["menu_consignacion"] == "modificar":
            st.subheader("Modificar Consignación Existente")
            try:
                con = conectar_bd()
                cur = con.cursor()
                cur.execute("SELECT numero_consignacion, nombre_consignante, rut_consignante, patente FROM consignaciones ORDER BY numero_consignacion DESC LIMIT 10")
                filas = cur.fetchall()
                if filas:
                    st.markdown("📋 **Últimas Consignaciones guardadas:**")
                    st.dataframe(pd.DataFrame(filas, columns=["N° Doc", "Nombre", "RUT", "Patente"]), use_container_width=True, hide_index=True)
            except: pass
            finally:
                if "cur" in locals(): cur.close()
                if "con" in locals(): con.close()

            st.write("---")
            nro_mod_c = st.text_input("Ingresa el N° a buscar:", key="nro_modificar_consig")
            if st.button("🔎 Buscar", key="btn_busc_con"):
                try:
                    con = conectar_bd()
                    cur = con.cursor()
                    cur.execute("SELECT * FROM consignaciones WHERE numero_consignacion = %s", (nro_mod_c.strip(),))
                    cols = [d[0] for d in cur.description]
                    fila = cur.fetchone()
                    if fila:
                        st.session_state["consig_mod"] = dict(zip(cols, fila))
                        st.success("✅ Encontrada.")
                    else: st.error("❌ No se encontró.")
                except Exception as e: st.error(f"Error: {e}")
                finally:
                    if "cur" in locals(): cur.close()
                    if "con" in locals(): con.close()

            if "consig_mod" in st.session_state:
                c = st.session_state["consig_mod"]
                with st.form("form_mod_consignacion"):
                    col1, col2 = st.columns(2)
                    with col1:
                        fecha_mc = st.text_input("Fecha", value=c.get("fecha",""))
                        nombre_mc = st.text_input("Nombre Consignante", value=c.get("nombre_consignante",""))
                        rut_mc = st.text_input("RUT Consignante", value=c.get("rut_consignante",""))
                        dir_mc = st.text_input("Dirección", value=c.get("direccion",""))
                        comuna_mc = st.text_input("Comuna", value=c.get("comuna",""))
                        tel_mc = st.text_input("Teléfono", value=c.get("telefono",""))
                        correo_mc = st.text_input("Correo", value=c.get("correo",""))
                        precio_mc = st.text_input("Precio de Venta", value=c.get("precio",""))
                        m_escrito_mc = st.text_input("Precio en Palabras", value=c.get("m_escrito",""))
                        comision_mc = st.text_input("Comisión %", value=c.get("comision",""))
                        d_contrato_mc = st.text_input("Duración Contrato", value=c.get("d_contrato",""))
                        pago_consignante_mc = st.text_input("Método Pago", value=c.get("pago_consignante",""))
                        cb_nombre_mc = st.text_input("Nombre Cuenta", value=c.get("cb_nombre",""))
                        cb_rut_mc = st.text_input("RUT Cuenta", value=c.get("cb_rut",""))
                        cb_banco_mc = st.text_input("Banco", value=c.get("cb_banco",""))
                        cb_cuenta_mc = st.text_input("Tipo Cuenta", value=c.get("cb_cuenta",""))
                        cb_numero_mc = st.text_input("Número Cuenta", value=c.get("cb_numero",""))
                    with col2:
                        tipo_veh_mc = st.text_input("Tipo Vehículo", value=c.get("tipo_vehiculo",""))
                        marca_mc = st.text_input("Marca", value=c.get("marca",""))
                        modelo_mc = st.text_input("Modelo", value=c.get("modelo",""))
                        anio_mc = st.text_input("Año", value=c.get("anio",""))
                        motor_mc = st.text_input("N° Motor", value=c.get("motor",""))
                        chasis_mc = st.text_input("N° Chasis", value=c.get("chasis",""))
                        color_mc = st.text_input("Color", value=c.get("color",""))
                        pat_mc = st.text_input("Patente", value=c.get("patente",""))
                        kms_mc = st.text_input("Kilometraje", value=c.get("kms",""))
                        obs_mc = st.text_area("Observaciones", value=c.get("observaciones",""))
                        
                    actualizar_y_generar_c = st.form_submit_button("💾 Actualizar y Reimprimir Word")

                if actualizar_y_generar_c:
                    v_consig, r_consig = validar_rut(rut_mc)

                    if not v_consig: st.error("❌ RUT Consignante inválido.")
                    else:
                        try:
                            con = conectar_bd()
                            cur = con.cursor()
                            cur.execute("""
                                UPDATE consignaciones SET
                                    fecha=%s, nombre_consignante=%s, rut_consignante=%s, direccion=%s, 
                                    comuna=%s, telefono=%s, correo=%s, precio=%s, m_escrito=%s, comision=%s, 
                                    d_contrato=%s, pago_consignante=%s, cb_nombre=%s, cb_rut=%s, cb_banco=%s, 
                                    cb_cuenta=%s, cb_numero=%s, tipo_vehiculo=%s, marca=%s, modelo=%s, anio=%s, 
                                    motor=%s, chasis=%s, color=%s, patente=%s, kms=%s, observaciones=%s
                                WHERE numero_consignacion=%s
                            """, (
                                fecha_mc, nombre_mc, r_consig, dir_mc, comuna_mc, tel_mc, correo_mc, precio_mc, 
                                m_escrito_mc, comision_mc, d_contrato_mc, pago_consignante_mc, cb_nombre_mc, cb_rut_mc, cb_banco_mc,
                                cb_cuenta_mc, cb_numero_mc, tipo_veh_mc, marca_mc, modelo_mc, anio_mc,
                                motor_mc, chasis_mc, color_mc, pat_mc, kms_mc, obs_mc, c["numero_consignacion"]
                            ))
                            con.commit()

                            ruta_plantilla = "Plantillas/Contrato_consignacion.docx"
                            doc = Document(ruta_plantilla)
                            reemplazos = {
                                "[NÚMERO_CONSIGNACION]": c["numero_consignacion"], "[FECHA]": fecha_mc,
                                "[NOMBRE_CONSIGNANTE]": nombre_mc, "[RUT_CONSIGNANTE]": r_consig,
                                "[DIRECCION]": dir_mc, "[COMUNA]": comuna_mc, "[TELEFONO]": tel_mc, 
                                "[CORREO]": correo_mc, "[PRECIO]": precio_mc, "[M_ESCRITO]": m_escrito_mc, 
                                "[COMISION]": comision_mc, "[D_CONTRATO]": d_contrato_mc, "[PAGO_CONSIGNANTE]": pago_consignante_mc,
                                "[CB_NOMBRE]": cb_nombre_mc, "[CB_RUT]": cb_rut_mc, "[CB_BANCO]": cb_banco_mc, 
                                "[CB_CUENTA]": cb_cuenta_mc, "[CB_NUMERO]": cb_numero_mc, "[TIPO_VEHICULO]": tipo_veh_mc,
                                "[MARCA]": marca_mc, "[MODELO]": modelo_mc, "[AÑO]": anio_mc, 
                                "[MOTOR]": motor_mc, "[CHASIS]": chasis_mc, "[COLOR]": color_mc, 
                                "[PATENTE]": pat_mc, "[KMS]": kms_mc, "[OBSERVACIONES]": obs_mc
                            }
                            for p in doc.paragraphs: reemplazar_en_parrafo(p, reemplazos)
                            for tabla in doc.tables:
                                for fila in tabla.rows:
                                    for celda in fila.cells:
                                        for p in celda.paragraphs: reemplazar_en_parrafo(p, reemplazos)
                            buffer = io.BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.success("✅ ¡Actualizado y reimpreso!")
                            st.download_button("📥 Descargar", data=buffer, file_name=f"Consignacion_{c['numero_consignacion']}_Modificada.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                            del st.session_state["consig_mod"]
                        except Exception as e: st.error(f"❌ Error: {e}")
                        finally:
                            if "cur" in locals(): cur.close()
                            if "con" in locals(): con.close()

    # ---------------------------------------------------------
    # MÓDULO 3: DECLARACIÓN JURADA
    # ---------------------------------------------------------
    with sub_declaracion:
        if "menu_declaracion" not in st.session_state:
            st.session_state["menu_declaracion"] = "crear"

        st.markdown("**¿Qué deseas hacer con la Declaración Jurada?**")
        col_btn5, col_btn6 = st.columns(2)
        with col_btn5:
            if st.button("➕ Crear y guardar", key="btn_dec_crear", use_container_width=True):
                st.session_state["menu_declaracion"] = "crear"
        with col_btn6:
            if st.button("✏️ Modificar / Reimprimir", key="btn_dec_mod", use_container_width=True):
                st.session_state["menu_declaracion"] = "modificar"
        st.write("")

        if st.session_state["menu_declaracion"] == "crear":
            st.subheader("Crear Nueva Declaración Jurada")
            with st.form("form_nueva_declaracion", clear_on_submit=False):
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("📝 **Datos de la Declaración**")
                    nro_decl_val = st.text_input("Número de Declaración *", placeholder="Ej: 001")
                    fecha_decl_val = st.text_input("Fecha")
                    st.markdown("👤 **Datos de la Compradora**")
                    nombre_comp_val = st.text_input("Nombre Compradora *")
                    rut_comp_val = st.text_input("RUT Compradora")
                    nac_comp_val = st.text_input("Nacionalidad")
                    dir_comp_val = st.text_input("Dirección")
                    com_comp_val = st.text_input("Comuna")
                    st.markdown("💰 **Datos del Trato**")
                    venta_v_val = st.text_input("Venta_V (Monto en Números)")
                    monto_escrito_val = st.text_input("Monto Escrito")
                with col2:
                    st.markdown("🚗 **Datos del Vehículo**")
                    vehiculo_val = st.text_input("Vehículo (Tipo)")
                    marca_d_val = st.text_input("Marca")
                    modelo_d_val = st.text_input("Modelo")
                    anio_d_val = st.text_input("Año")
                    motor_d_val = st.text_input("N° Motor")
                    chasis_d_val = st.text_input("N° Chasis / VIN")
                    color_d_val = st.text_input("Color")
                    patente_d_val = st.text_input("Patente")

                generar_decl = st.form_submit_button("💾 Guardar y Generar")

            if generar_decl:
                if not nro_decl_val or not nombre_comp_val:
                    st.error("⚠️ El Número y Nombre son obligatorios.")
                else:
                    v_comp, r_comp = validar_rut(rut_comp_val)
                    if not v_comp: st.error("❌ RUT Compradora inválido.")
                    else:
                        try:
                            con = conectar_bd()
                            cur = con.cursor()
                            cur.execute("""
                                INSERT INTO declaraciones (
                                    numero_declaracion, fecha, nombre_compradora, rut_compradora, nac_compradora, 
                                    direccion_compradora, comuna_compradora, vehiculo, marca, modelo, anio, 
                                    motor, chasis, color, patente, venta_v, monto_escrito
                                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                            """, (
                                nro_decl_val, fecha_decl_val, nombre_comp_val, r_comp, nac_comp_val,
                                dir_comp_val, com_comp_val, vehiculo_val, marca_d_val, modelo_d_val, anio_d_val,
                                motor_d_val, chasis_d_val, color_d_val, patente_d_val, venta_v_val, monto_escrito_val
                            ))
                            con.commit()

                            ruta_plantilla = "Plantillas/Declaracion_Jurada.docx"
                            doc = Document(ruta_plantilla)
                            reemplazos = {
                                "[NÚMERO_DECLARACION]": nro_decl_val, "[FECHA]": fecha_decl_val,
                                "[NOMBRE_COMPRADORA]": nombre_comp_val, "[RUT_COMPRADORA]": r_comp,
                                "[NAC_COMPRADORA]": nac_comp_val, "[DIRECCION_COMPRADORA]": dir_comp_val,
                                "[COMUNA_COMPRADORA]": com_comp_val, "[VEHICULO]": vehiculo_val,
                                "[MARCA]": marca_d_val, "[MODELO]": modelo_d_val, "[AÑO]": anio_d_val, 
                                "[N_MOTOR]": motor_d_val, "[N_CHASIS]": chasis_d_val, "[COLOR]": color_d_val, 
                                "[PATENTE]": patente_d_val, "[VENTA_V]": venta_v_val, "[MONTO_ESCRITO]": monto_escrito_val
                            }
                            for p in doc.paragraphs: reemplazar_en_parrafo(p, reemplazos)
                            for tabla in doc.tables:
                                for fila in tabla.rows:
                                    for celda in fila.cells:
                                        for p in celda.paragraphs: reemplazar_en_parrafo(p, reemplazos)
                            buffer = io.BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.success(f"✅ ¡Guardado y generado!")
                            st.download_button("📥 Descargar", data=buffer, file_name=f"Declaracion_{nro_decl_val}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        except pg_errors.UniqueViolation: st.error("❌ Ese Número ya existe.")
                        except Exception as e: st.error(f"❌ Error: {e}")
                        finally:
                            if "cur" in locals(): cur.close()
                            if "con" in locals(): con.close()

        elif st.session_state["menu_declaracion"] == "modificar":
            st.subheader("Modificar Declaración Jurada Existente")
            try:
                con = conectar_bd()
                cur = con.cursor()
                cur.execute("SELECT numero_declaracion, nombre_compradora, rut_compradora, patente FROM declaraciones ORDER BY numero_declaracion DESC LIMIT 10")
                filas = cur.fetchall()
                if filas:
                    st.markdown("📋 **Últimas Declaraciones:**")
                    st.dataframe(pd.DataFrame(filas, columns=["N° Doc", "Nombre", "RUT", "Patente"]), use_container_width=True, hide_index=True)
            except: pass
            finally:
                if "cur" in locals(): cur.close()
                if "con" in locals(): con.close()

            st.write("---")
            nro_mod_d = st.text_input("Ingresa el N° a buscar:", key="nro_mod_decl")
            if st.button("🔎 Buscar", key="btn_busc_dec"):
                try:
                    con = conectar_bd()
                    cur = con.cursor()
                    cur.execute("SELECT * FROM declaraciones WHERE numero_declaracion = %s", (nro_mod_d.strip(),))
                    cols = [d[0] for d in cur.description]
                    fila = cur.fetchone()
                    if fila:
                        st.session_state["decl_mod"] = dict(zip(cols, fila))
                        st.success("✅ Encontrada.")
                    else: st.error("❌ No se encontró.")
                except Exception as e: st.error(f"Error: {e}")
                finally:
                    if "cur" in locals(): cur.close()
                    if "con" in locals(): con.close()

            if "decl_mod" in st.session_state:
                d = st.session_state["decl_mod"]
                with st.form("form_mod_declaracion"):
                    col1, col2 = st.columns(2)
                    with col1:
                        fecha_md = st.text_input("Fecha", value=d.get("fecha",""))
                        nombre_comp_md = st.text_input("Nombre Compradora", value=d.get("nombre_compradora",""))
                        rut_comp_md = st.text_input("RUT Compradora", value=d.get("rut_compradora",""))
                        nac_comp_md = st.text_input("Nacionalidad", value=d.get("nac_compradora",""))
                        dir_comp_md = st.text_input("Dirección", value=d.get("direccion_compradora",""))
                        com_comp_md = st.text_input("Comuna", value=d.get("comuna_compradora",""))
                        venta_v_md = st.text_input("Venta_V", value=d.get("venta_v",""))
                        monto_escrito_md = st.text_input("Monto Escrito", value=d.get("monto_escrito",""))
                    with col2:
                        vehiculo_md = st.text_input("Vehículo (Tipo)", value=d.get("vehiculo",""))
                        marca_md = st.text_input("Marca", value=d.get("marca",""))
                        modelo_md = st.text_input("Modelo", value=d.get("modelo",""))
                        anio_md = st.text_input("Año", value=d.get("anio",""))
                        motor_md = st.text_input("N° Motor", value=d.get("motor",""))
                        chasis_md = st.text_input("N° Chasis", value=d.get("chasis",""))
                        color_md = st.text_input("Color", value=d.get("color",""))
                        patente_md = st.text_input("Patente", value=d.get("patente",""))
                    actualizar_y_generar_d = st.form_submit_button("💾 Actualizar y Reimprimir Word")

                if actualizar_y_generar_d:
                    v_comp, r_comp = validar_rut(rut_comp_md)
                    if not v_comp: st.error("❌ RUT Compradora inválido.")
                    else:
                        try:
                            con = conectar_bd()
                            cur = con.cursor()
                            cur.execute("""
                                UPDATE declaraciones SET
                                    fecha=%s, nombre_compradora=%s, rut_compradora=%s, nac_compradora=%s, 
                                    direccion_compradora=%s, comuna_compradora=%s, vehiculo=%s, marca=%s, 
                                    modelo=%s, anio=%s, motor=%s, chasis=%s, color=%s, patente=%s, 
                                    venta_v=%s, monto_escrito=%s
                                WHERE numero_declaracion=%s
                            """, (
                                fecha_md, nombre_comp_md, r_comp, nac_comp_md, dir_comp_md, com_comp_md, 
                                vehiculo_md, marca_md, modelo_md, anio_md, motor_md, chasis_md, color_md, 
                                patente_md, venta_v_md, monto_escrito_md, d["numero_declaracion"]
                            ))
                            con.commit()

                            ruta_plantilla = "Plantillas/Declaracion_Jurada.docx"
                            doc = Document(ruta_plantilla)
                            reemplazos = {
                                "[NÚMERO_DECLARACION]": d["numero_declaracion"], "[FECHA]": fecha_md,
                                "[NOMBRE_COMPRADORA]": nombre_comp_md, "[RUT_COMPRADORA]": r_comp,
                                "[NAC_COMPRADORA]": nac_comp_md, "[DIRECCION_COMPRADORA]": dir_comp_md,
                                "[COMUNA_COMPRADORA]": com_comp_md, "[VEHICULO]": vehiculo_md,
                                "[MARCA]": marca_md, "[MODELO]": modelo_md, "[AÑO]": anio_md, 
                                "[N_MOTOR]": motor_md, "[N_CHASIS]": chasis_md, "[COLOR]": color_md, 
                                "[PATENTE]": patente_md, "[VENTA_V]": venta_v_md, "[MONTO_ESCRITO]": monto_escrito_md
                            }
                            for p in doc.paragraphs: reemplazar_en_parrafo(p, reemplazos)
                            for tabla in doc.tables:
                                for fila in tabla.rows:
                                    for celda in fila.cells:
                                        for p in celda.paragraphs: reemplazar_en_parrafo(p, reemplazos)
                            buffer = io.BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.success("✅ ¡Actualizado y reimpreso!")
                            st.download_button("📥 Descargar", data=buffer, file_name=f"Declaracion_{d['numero_declaracion']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                            del st.session_state["decl_mod"]
                        except Exception as e: st.error(f"❌ Error: {e}")
                        finally:
                            if "cur" in locals(): cur.close()
                            if "con" in locals(): con.close()

    # ---------------------------------------------------------
    # MÓDULO 4: CARTA PODER
    # ---------------------------------------------------------
    with sub_poder:
        if "menu_poder" not in st.session_state:
            st.session_state["menu_poder"] = "crear"

        st.markdown("**¿Qué deseas hacer con la Carta Poder?**")
        col_btn7, col_btn8 = st.columns(2)
        with col_btn7:
            if st.button("➕ Crear y guardar", key="btn_pod_crear", use_container_width=True):
                st.session_state["menu_poder"] = "crear"
        with col_btn8:
            if st.button("✏️ Modificar / Reimprimir", key="btn_pod_mod", use_container_width=True):
                st.session_state["menu_poder"] = "modificar"
        st.write("")

        if st.session_state["menu_poder"] == "crear":
            st.subheader("Crear Nueva Carta Poder")
            with st.form("form_nueva_carta_poder", clear_on_submit=False):
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("📝 **Datos del Mandante (Quien da el poder)**")
                    nro_poder_val = st.text_input("Número de Poder *")
                    fecha_p_val = st.text_input("Fecha")
                    nombre_p_val = st.text_input("Nombre de quien da el poder *")
                    rut_p_val = st.text_input("RUT de quien da el poder")
                    dir_p_val = st.text_input("Dirección")
                    com_p_val = st.text_input("Comuna")
                    st.markdown("👤 **Datos del Representante (Quien recibe el poder)**")
                    rep_p_val = st.text_input("Representante")
                    rut_rep_p_val = st.text_input("RUT Representante")
                    duracion_p_val = st.text_input("Duración del Poder")
                with col2:
                    st.markdown("🚗 **Datos del Vehículo**")
                    tipo_veh_p_val = st.text_input("Tipo de Vehículo")
                    marca_p_val = st.text_input("Marca")
                    modelo_p_val = st.text_input("Modelo")
                    anio_p_val = st.text_input("Año")
                    motor_p_val = st.text_input("N° Motor")
                    chasis_p_val = st.text_input("N° Chasis")
                    vin_p_val = st.text_input("N° VIN")
                    color_p_val = st.text_input("Color")
                    pat_p_val = st.text_input("Patente")

                generar_poder = st.form_submit_button("💾 Guardar y Generar")

            if generar_poder:
                if not nro_poder_val or not nombre_p_val:
                    st.error("⚠️ El Número y Nombre son obligatorios.")
                else:
                    v_pod, r_pod = validar_rut(rut_p_val)
                    v_rep, r_rep = validar_rut(rut_rep_p_val)

                    if not v_pod: st.error("❌ RUT Mandante inválido.")
                    elif not v_rep: st.error("❌ RUT Representante inválido.")
                    else:
                        try:
                            con = conectar_bd()
                            cur = con.cursor()
                            cur.execute("""
                                INSERT INTO cartas_poder (
                                    numero_poder, fecha, nombre_poder, rut_poder, representante_poder, 
                                    rut_poder_r, direccion_poder, comuna_poder, tipo_vehiculo, marca, 
                                    modelo, anio, motor, chasis, n_vin, color, patente, duracion_poder
                                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                            """, (
                                nro_poder_val, fecha_p_val, nombre_p_val, r_pod, rep_p_val,
                                r_rep, dir_p_val, com_p_val, tipo_veh_p_val, marca_p_val,
                                modelo_p_val, anio_p_val, motor_p_val, chasis_p_val, vin_p_val,
                                color_p_val, pat_p_val, duracion_p_val
                            ))
                            con.commit()

                            ruta_plantilla = "Plantillas/Carta_Poder.docx"
                            doc = Document(ruta_plantilla)
                            reemplazos = {
                                "[NÚMERO_PODER]": nro_poder_val, "[FECHA]": fecha_p_val, "[NOMBRE_PODER]": nombre_p_val, 
                                "[RUT_PODER]": r_pod, "[REPRESENTANTE_PODER]": rep_p_val, "[RUT_PODER_R]": r_rep,
                                "[DIRECCION_PODER]": dir_p_val, "[COMUNA_PODER]": com_p_val, "[TIPO_VEHICULO]": tipo_veh_p_val, 
                                "[MARCA]": marca_p_val, "[MODELO]": modelo_p_val, "[AÑO]": anio_p_val, "[N_MOTOR]": motor_p_val, 
                                "[N_CHASIS]": chasis_p_val, "[N_VIN]": vin_p_val, "[COLOR]": color_p_val, "[PATENTE]": pat_p_val, 
                                "[DURACION_PODER]": duracion_p_val
                            }
                            for p in doc.paragraphs: reemplazar_en_parrafo(p, reemplazos)
                            for tabla in doc.tables:
                                for fila in tabla.rows:
                                    for celda in fila.cells:
                                        for p in celda.paragraphs: reemplazar_en_parrafo(p, reemplazos)
                            buffer = io.BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.success(f"✅ ¡Guardado y generado!")
                            st.download_button("📥 Descargar", data=buffer, file_name=f"Carta_Poder_{nro_poder_val}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        except pg_errors.UniqueViolation: st.error("❌ Ese Número ya existe.")
                        except Exception as e: st.error(f"❌ Error: {e}")
                        finally:
                            if "cur" in locals(): cur.close()
                            if "con" in locals(): con.close()

        elif st.session_state["menu_poder"] == "modificar":
            st.subheader("Modificar Carta Poder Existente")
            try:
                con = conectar_bd()
                cur = con.cursor()
                cur.execute("SELECT numero_poder, nombre_poder, rut_poder, patente FROM cartas_poder ORDER BY numero_poder DESC LIMIT 10")
                filas = cur.fetchall()
                if filas:
                    st.markdown("📋 **Últimos Poderes:**")
                    st.dataframe(pd.DataFrame(filas, columns=["N° Doc", "Nombre", "RUT", "Patente"]), use_container_width=True, hide_index=True)
            except: pass
            finally:
                if "cur" in locals(): cur.close()
                if "con" in locals(): con.close()

            st.write("---")
            nro_mod_p = st.text_input("Ingresa el N° a buscar:", key="nro_mod_poder")
            if st.button("🔎 Buscar", key="btn_busc_pod"):
                try:
                    con = conectar_bd()
                    cur = con.cursor()
                    cur.execute("SELECT * FROM cartas_poder WHERE numero_poder = %s", (nro_mod_p.strip(),))
                    cols = [d[0] for d in cur.description]
                    fila = cur.fetchone()
                    if fila:
                        st.session_state["poder_mod"] = dict(zip(cols, fila))
                        st.success("✅ Encontrada.")
                    else: st.error("❌ No se encontró.")
                except Exception as e: st.error(f"Error: {e}")
                finally:
                    if "cur" in locals(): cur.close()
                    if "con" in locals(): con.close()

            if "poder_mod" in st.session_state:
                p = st.session_state["poder_mod"]
                with st.form("form_mod_poder"):
                    col1, col2 = st.columns(2)
                    with col1:
                        fecha_mp = st.text_input("Fecha", value=p.get("fecha",""))
                        nombre_p_mp = st.text_input("Nombre de quien da el poder", value=p.get("nombre_poder",""))
                        rut_p_mp = st.text_input("RUT de quien da el poder", value=p.get("rut_poder",""))
                        dir_p_mp = st.text_input("Dirección", value=p.get("direccion_poder",""))
                        com_p_mp = st.text_input("Comuna", value=p.get("comuna_poder",""))
                        rep_p_mp = st.text_input("Representante", value=p.get("representante_poder",""))
                        rut_rep_p_mp = st.text_input("RUT Representante", value=p.get("rut_poder_r",""))
                        duracion_p_mp = st.text_input("Duración del Poder", value=p.get("duracion_poder",""))
                    with col2:
                        tipo_veh_p_mp = st.text_input("Tipo de Vehículo", value=p.get("tipo_vehiculo",""))
                        marca_p_mp = st.text_input("Marca", value=p.get("marca",""))
                        modelo_p_mp = st.text_input("Modelo", value=p.get("modelo",""))
                        anio_p_mp = st.text_input("Año", value=p.get("anio",""))
                        motor_p_mp = st.text_input("N° Motor", value=p.get("motor",""))
                        chasis_p_mp = st.text_input("N° Chasis", value=p.get("chasis",""))
                        vin_p_mp = st.text_input("N° VIN", value=p.get("n_vin",""))
                        color_p_mp = st.text_input("Color", value=p.get("color",""))
                        pat_p_mp = st.text_input("Patente", value=p.get("patente",""))
                    actualizar_y_generar_p = st.form_submit_button("💾 Actualizar y Reimprimir Word")

                if actualizar_y_generar_p:
                    v_pod, r_pod = validar_rut(rut_p_mp)
                    v_rep, r_rep = validar_rut(rut_rep_p_mp)

                    if not v_pod: st.error("❌ RUT Mandante inválido.")
                    elif not v_rep: st.error("❌ RUT Representante inválido.")
                    else:
                        try:
                            con = conectar_bd()
                            cur = con.cursor()
                            cur.execute("""
                                UPDATE cartas_poder SET
                                    fecha=%s, nombre_poder=%s, rut_poder=%s, representante_poder=%s, 
                                    rut_poder_r=%s, direccion_poder=%s, comuna_poder=%s, tipo_vehiculo=%s, 
                                    marca=%s, modelo=%s, anio=%s, motor=%s, chasis=%s, n_vin=%s, color=%s, 
                                    patente=%s, duracion_poder=%s
                                WHERE numero_poder=%s
                            """, (
                                fecha_mp, nombre_p_mp, r_pod, rep_p_mp, r_rep, dir_p_mp, com_p_mp, 
                                tipo_veh_p_mp, marca_p_mp, modelo_p_mp, anio_p_mp, motor_p_mp, chasis_p_mp, 
                                vin_p_mp, color_p_mp, pat_p_mp, duracion_p_mp, p["numero_poder"]
                            ))
                            con.commit()
                            ruta_plantilla = "Plantillas/Carta_Poder.docx"
                            doc = Document(ruta_plantilla)
                            reemplazos = {
                                "[NÚMERO_PODER]": p["numero_poder"], "[FECHA]": fecha_mp, "[NOMBRE_PODER]": nombre_p_mp, 
                                "[RUT_PODER]": r_pod, "[REPRESENTANTE_PODER]": rep_p_mp, "[RUT_PODER_R]": r_rep,
                                "[DIRECCION_PODER]": dir_p_mp, "[COMUNA_PODER]": com_p_mp, "[TIPO_VEHICULO]": tipo_veh_p_mp, 
                                "[MARCA]": marca_p_mp, "[MODELO]": modelo_p_mp, "[AÑO]": anio_p_mp, "[N_MOTOR]": motor_p_mp, 
                                "[N_CHASIS]": chasis_p_mp, "[N_VIN]": vin_p_mp, "[COLOR]": color_p_mp, "[PATENTE]": pat_p_mp, 
                                "[DURACION_PODER]": duracion_p_mp
                            }
                            for pr in doc.paragraphs: reemplazar_en_parrafo(pr, reemplazos)
                            for tabla in doc.tables:
                                for fila in tabla.rows:
                                    for celda in fila.cells:
                                        for pr in celda.paragraphs: reemplazar_en_parrafo(pr, reemplazos)
                            buffer = io.BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.success("✅ ¡Actualizado y reimpreso!")
                            st.download_button("📥 Descargar", data=buffer, file_name=f"Carta_Poder_{p['numero_poder']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                            del st.session_state["poder_mod"]
                        except Exception as e: st.error(f"❌ Error: {e}")
                        finally:
                            if "cur" in locals(): cur.close()
                            if "con" in locals(): con.close()

    # ---------------------------------------------------------
    # MÓDULO 5: CONTRATO NOTARIAL
    # ---------------------------------------------------------
    with sub_notarial:
        if "menu_notarial" not in st.session_state:
            st.session_state["menu_notarial"] = "crear"

        st.markdown("**¿Qué deseas hacer con el Contrato Notarial?**")
        col_btn9, col_btn10 = st.columns(2)
        with col_btn9:
            if st.button("➕ Crear y guardar", key="btn_not_crear", use_container_width=True):
                st.session_state["menu_notarial"] = "crear"
        with col_btn10:
            if st.button("✏️ Modificar / Reimprimir", key="btn_not_mod", use_container_width=True):
                st.session_state["menu_notarial"] = "modificar"
        st.write("")

        if st.session_state["menu_notarial"] == "crear":
            st.subheader("Crear Nuevo Contrato Notarial")
            with st.form("form_nuevo_contrato_notarial", clear_on_submit=False):
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("📝 **Datos de la Venta y Comprador**")
                    nro_repertorio_val = st.text_input("Número de Repertorio *")
                    fecha_notarial_val = st.text_input("Fecha")
                    nombre_comp_not_val = st.text_input("Nombre del Comprador *")
                    rut_comp_not_val = st.text_input("RUT del Comprador")
                    nac_comp_not_val = st.text_input("Nacionalidad")
                    dir_comp_not_val = st.text_input("Dirección")
                    com_comp_not_val = st.text_input("Comuna")
                    st.markdown("💰 **Datos Financieros**")
                    m_venta_not_val = st.text_input("Monto Venta")
                    monto_escrito_not_val = st.text_input("Monto Escrito")
                with col2:
                    st.markdown("🚗 **Datos del Vehículo**")
                    tipo_veh_not_val = st.text_input("Tipo de Vehículo")
                    marca_not_val = st.text_input("Marca")
                    modelo_not_val = st.text_input("Modelo")
                    anio_not_val = st.text_input("Año")
                    motor_not_val = st.text_input("N° Motor")
                    chasis_not_val = st.text_input("N° Chasis")
                    vin_not_val = st.text_input("N° VIN")
                    color_not_val = st.text_input("Color")
                    patente_not_val = st.text_input("Patente")

                generar_notarial = st.form_submit_button("💾 Guardar y Generar")

            if generar_notarial:
                if not nro_repertorio_val or not nombre_comp_not_val:
                    st.error("⚠️ El Número y Nombre son obligatorios.")
                else:
                    v_comp, r_comp = validar_rut(rut_comp_not_val)
                    if not v_comp: st.error("❌ RUT Comprador inválido.")
                    else:
                        try:
                            con = conectar_bd()
                            cur = con.cursor()
                            cur.execute("""
                                INSERT INTO contratos_notariales (
                                    numero_repertorio, fecha, nombre_comprador, rut_comprador, nac_comprador, 
                                    direccion_comprador, comuna_comprador, tipo_vehiculo, marca, modelo, 
                                    anio, motor, chasis, n_vin, color, patente, m_venta, monto_escrito
                                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                            """, (
                                nro_repertorio_val, fecha_notarial_val, nombre_comp_not_val, r_comp, 
                                nac_comp_not_val, dir_comp_not_val, com_comp_not_val, tipo_veh_not_val, 
                                marca_not_val, modelo_not_val, anio_not_val, motor_not_val, chasis_not_val, 
                                vin_not_val, color_not_val, patente_not_val, m_venta_not_val, monto_escrito_not_val
                            ))
                            con.commit()

                            ruta_plantilla = "Plantillas/Contrato_Notarial.docx"
                            doc = Document(ruta_plantilla)
                            reemplazos = {
                                "[NUMERO_REPERTORIO]": nro_repertorio_val, "[FECHA]": fecha_notarial_val,
                                "[NOMBRE_COMPRADOR]": nombre_comp_not_val, "[RUT_COMPRADOR]": r_comp,
                                "[NAC_COMPRADOR]": nac_comp_not_val, "[DIRECCION_COMPRADOR]": dir_comp_not_val,
                                "[COMUNA_COMPRADOR]": com_comp_not_val, "[TIPO_VEHICULO]": tipo_veh_not_val,
                                "[MARCA]": marca_not_val, "[MODELO]": modelo_not_val, "[AÑO]": anio_not_val,
                                "[N_MOTOR]": motor_not_val, "[N_CHASIS]": chasis_not_val, "[N_VIN]": vin_not_val, 
                                "[COLOR]": color_not_val, "[PATENTE]": patente_not_val, "[M_VENTA]": m_venta_not_val, 
                                "[MONTO_ESCRITO]": monto_escrito_not_val
                            }
                            for p in doc.paragraphs: reemplazar_en_parrafo(p, reemplazos)
                            for tabla in doc.tables:
                                for fila in tabla.rows:
                                    for celda in fila.cells:
                                        for p in celda.paragraphs: reemplazar_en_parrafo(p, reemplazos)
                            buffer = io.BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.success(f"✅ ¡Guardado y generado!")
                            st.download_button("📥 Descargar", data=buffer, file_name=f"Notarial_{nro_repertorio_val}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        except pg_errors.UniqueViolation: st.error("❌ Ese Número ya existe.")
                        except Exception as e: st.error(f"❌ Error: {e}")
                        finally:
                            if "cur" in locals(): cur.close()
                            if "con" in locals(): con.close()

        elif st.session_state["menu_notarial"] == "modificar":
            st.subheader("Modificar Contrato Notarial Existente")
            try:
                con = conectar_bd()
                cur = con.cursor()
                cur.execute("SELECT numero_repertorio, nombre_comprador, rut_comprador, patente FROM contratos_notariales ORDER BY numero_repertorio DESC LIMIT 10")
                filas = cur.fetchall()
                if filas:
                    st.markdown("📋 **Últimos Contratos Notariales:**")
                    st.dataframe(pd.DataFrame(filas, columns=["N° Doc", "Nombre", "RUT", "Patente"]), use_container_width=True, hide_index=True)
            except: pass
            finally:
                if "cur" in locals(): cur.close()
                if "con" in locals(): con.close()

            st.write("---")
            nro_mod_not = st.text_input("Ingresa el N° a buscar:", key="nro_mod_notarial")
            if st.button("🔎 Buscar", key="btn_busc_not"):
                try:
                    con = conectar_bd()
                    cur = con.cursor()
                    cur.execute("SELECT * FROM contratos_notariales WHERE numero_repertorio = %s", (nro_mod_not.strip(),))
                    cols = [d[0] for d in cur.description]
                    fila = cur.fetchone()
                    if fila:
                        st.session_state["notarial_mod"] = dict(zip(cols, fila))
                        st.success("✅ Encontrado.")
                    else: st.error("❌ No se encontró.")
                except Exception as e: st.error(f"Error: {e}")
                finally:
                    if "cur" in locals(): cur.close()
                    if "con" in locals(): con.close()

            if "notarial_mod" in st.session_state:
                n_mod = st.session_state["notarial_mod"]
                with st.form("form_mod_notarial"):
                    col1, col2 = st.columns(2)
                    with col1:
                        fecha_mn = st.text_input("Fecha", value=n_mod.get("fecha",""))
                        nombre_comp_mn = st.text_input("Nombre Comprador", value=n_mod.get("nombre_comprador",""))
                        rut_comp_mn = st.text_input("RUT Comprador", value=n_mod.get("rut_comprador",""))
                        nac_comp_mn = st.text_input("Nacionalidad", value=n_mod.get("nac_comprador",""))
                        dir_comp_mn = st.text_input("Dirección", value=n_mod.get("direccion_comprador",""))
                        com_comp_mn = st.text_input("Comuna", value=n_mod.get("comuna_comprador",""))
                        m_venta_mn = st.text_input("Monto Venta", value=n_mod.get("m_venta",""))
                        monto_escrito_mn = st.text_input("Monto Escrito", value=n_mod.get("monto_escrito",""))
                    with col2:
                        tipo_veh_mn = st.text_input("Tipo Vehículo", value=n_mod.get("tipo_vehiculo",""))
                        marca_mn = st.text_input("Marca", value=n_mod.get("marca",""))
                        modelo_mn = st.text_input("Modelo", value=n_mod.get("modelo",""))
                        anio_mn = st.text_input("Año", value=n_mod.get("anio",""))
                        motor_mn = st.text_input("N° Motor", value=n_mod.get("motor",""))
                        chasis_mn = st.text_input("N° Chasis", value=n_mod.get("chasis",""))
                        vin_mn = st.text_input("N° VIN", value=n_mod.get("n_vin",""))
                        color_mn = st.text_input("Color", value=n_mod.get("color",""))
                        patente_mn = st.text_input("Patente", value=n_mod.get("patente",""))
                    actualizar_y_generar_n = st.form_submit_button("💾 Actualizar y Reimprimir Word")

                if actualizar_y_generar_n:
                    v_comp, r_comp = validar_rut(rut_comp_mn)
                    if not v_comp: st.error("❌ RUT Comprador inválido.")
                    else:
                        try:
                            con = conectar_bd()
                            cur = con.cursor()
                            cur.execute("""
                                UPDATE contratos_notariales SET
                                    fecha=%s, nombre_comprador=%s, rut_comprador=%s, nac_comprador=%s, 
                                    direccion_comprador=%s, comuna_comprador=%s, tipo_vehiculo=%s, marca=%s, 
                                    modelo=%s, anio=%s, motor=%s, chasis=%s, n_vin=%s, color=%s, patente=%s, 
                                    m_venta=%s, monto_escrito=%s
                                WHERE numero_repertorio=%s
                            """, (
                                fecha_mn, nombre_comp_mn, r_comp, nac_comp_mn, dir_comp_mn, com_comp_mn, 
                                tipo_veh_mn, marca_mn, modelo_mn, anio_mn, motor_mn, chasis_mn, vin_mn, 
                                color_mn, patente_mn, m_venta_mn, monto_escrito_mn, n_mod["numero_repertorio"]
                            ))
                            con.commit()
                            ruta_plantilla = "Plantillas/Contrato_Notarial.docx"
                            doc = Document(ruta_plantilla)
                            reemplazos = {
                                "[NUMERO_REPERTORIO]": n_mod["numero_repertorio"], "[FECHA]": fecha_mn,
                                "[NOMBRE_COMPRADOR]": nombre_comp_mn, "[RUT_COMPRADOR]": r_comp,
                                "[NAC_COMPRADOR]": nac_comp_mn, "[DIRECCION_COMPRADOR]": dir_comp_mn,
                                "[COMUNA_COMPRADOR]": com_comp_mn, "[TIPO_VEHICULO]": tipo_veh_mn,
                                "[MARCA]": marca_mn, "[MODELO]": modelo_mn, "[AÑO]": anio_mn, 
                                "[N_MOTOR]": motor_mn, "[N_CHASIS]": chasis_mn, "[N_VIN]": vin_mn, 
                                "[COLOR]": color_mn, "[PATENTE]": patente_mn, "[M_VENTA]": m_venta_mn, 
                                "[MONTO_ESCRITO]": monto_escrito_mn
                            }
                            for pr in doc.paragraphs: reemplazar_en_parrafo(pr, reemplazos)
                            for tabla in doc.tables:
                                for fila in tabla.rows:
                                    for celda in fila.cells:
                                        for pr in celda.paragraphs: reemplazar_en_parrafo(pr, reemplazos)
                            buffer = io.BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.success("✅ ¡Actualizado y reimpreso!")
                            st.download_button("📥 Descargar", data=buffer, file_name=f"Notarial_{n_mod['numero_repertorio']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                            del st.session_state["notarial_mod"]
                        except Exception as e: st.error(f"❌ Error: {e}")
                        finally:
                            if "cur" in locals(): cur.close()
                            if "con" in locals(): con.close()

if "documentos" in tabs:
    with tabs["documentos"]:
        render_documentos()

# ================================================================
# PESTAÑA 2 — GESTIÓN DE PRODUCTOS
# ================================================================
def obtener_siguiente_codigo_producto():
    """Calcula el próximo código correlativo (001, 002, 003...) en base al
    mayor código numérico ya registrado. Sin tope: al pasar de 999 sigue
    con 1000, 1001, etc."""
    try:
        con = conectar_bd()
        cur = con.cursor()
        cur.execute("SELECT MAX(CAST(codigo_producto AS INTEGER)) FROM productos WHERE codigo_producto ~ '^[0-9]+$'")
        maximo = cur.fetchone()[0]
        return f"{(maximo or 0) + 1:03d}"
    except Exception:
        return "001"
    finally:
        if "cur" in locals(): cur.close()
        if "con" in locals(): con.close()

@st.fragment
def render_productos():
    if "menu_productos" not in st.session_state:
        st.session_state["menu_productos"] = "crear"

    st.markdown("**¿Qué deseas hacer en Productos?**")
    col_pr1, col_pr2 = st.columns(2)
    with col_pr1:
        if st.button("➕ Registrar nuevo producto", key="btn_pr_crear", use_container_width=True):
            st.session_state["menu_productos"] = "crear"
    with col_pr2:
        if st.button("✏️ Modificar producto existente", key="btn_pr_mod", use_container_width=True):
            st.session_state["menu_productos"] = "modificar"

    if st.session_state.get("_flash_productos"):
        st.success(st.session_state.pop("_flash_productos"))
    st.write("")

    # --------------------------------------------------------
    # MÓDULO A: REGISTRAR NUEVO PRODUCTO
    # --------------------------------------------------------
    if st.session_state["menu_productos"] == "crear":
        st.subheader("Registrar Nuevo Producto")

        try:
            con = conectar_bd()
            cur = con.cursor()
            cur.execute("SELECT codigo_producto, descripcion, stock_actual FROM productos ORDER BY id DESC LIMIT 15")
            filas_prod = cur.fetchall()
            if filas_prod:
                st.markdown("📋 **Últimos Códigos de Producto registrados:**")
                st.dataframe(pd.DataFrame(filas_prod, columns=["Código", "Descripción", "Stock"]), use_container_width=True, hide_index=True)
        except: pass
        finally:
            if "cur" in locals(): cur.close()
            if "con" in locals(): con.close()

        gen_p = st.session_state.get("_gen_producto", 0)
        codigo_producto_p = obtener_siguiente_codigo_producto()

        st.markdown("**📦 Datos del Producto**")
        cp1, cp2 = st.columns(2)
        with cp1:
            st.text_input("Código Producto (autogenerado)", value=codigo_producto_p, disabled=True, key=f"codigo_prod_crear_{gen_p}")
            codigo_barra_p = st.text_input("Código de Barra", key=f"codigo_barra_crear_{gen_p}")
        with cp2:
            descripcion_p = st.text_input("Descripción *", key=f"descripcion_prod_crear_{gen_p}")

        st.markdown("---")
        st.markdown("### 📊 Stock")
        se1, se2, se3, se4 = st.columns(4)
        with se1:
            stock_actual_p = st.number_input("Stock Actual", min_value=0, step=1, key=f"stock_actual_crear_{gen_p}")
        with se2:
            stock_minimo_p = st.number_input("Stock Mínimo", min_value=0, step=1, key=f"stock_minimo_crear_{gen_p}")
        with se3:
            stock_reposicion_p = st.number_input(
                "Stock Reposición", min_value=0, step=1, key=f"stock_reposicion_crear_{gen_p}",
                help="Umbral manual: al llegar aquí, es momento de empezar a comprar.",
            )
        with se4:
            stock_maximo_p = st.number_input("Stock Máximo", min_value=0, step=1, key=f"stock_maximo_crear_{gen_p}")

        st.markdown("---")
        st.markdown("### 💰 Precio de Venta")
        st.markdown("_Puedes ingresar el costo con o sin puntos (Ej: 10.000 o 10000)_")
        pe1, pe2, pe3, pe4 = st.columns(4)
        with pe1:
            costo_key_p = f"costo_prod_crear_{gen_p}"
            txt_costo_p = st.text_input(
                "Costo ($) *", value="0", key=costo_key_p,
                on_change=formatear_monto_callback, args=(costo_key_p,),
            )
            costo_p = int(txt_costo_p.replace(".", "").replace("$", "").strip() or 0)
            st.caption(f"🔎 Sombra: **${costo_p:,.0f}**".replace(",", "."))
        with pe2:
            margen_p = st.number_input("Margen (%)", min_value=0.0, step=0.5, format="%.2f", key=f"margen_prod_crear_{gen_p}")
        with pe3:
            neto_p = round(costo_p * (1 + margen_p / 100))
            st.metric("Neto Afecto a IVA", f"${neto_p:,.0f}".replace(",", "."))
        with pe4:
            precio_iva_p = round(neto_p * (1 + IVA_RATE))
            st.metric("Precio con IVA", f"${precio_iva_p:,.0f}".replace(",", "."))

        st.write("")
        guardar_producto = st.button("💾 Guardar Producto", use_container_width=True, type="primary")

        if guardar_producto:
            if not descripcion_p.strip():
                st.error("⚠️ La Descripción es obligatoria.")
            else:
                try:
                    con = conectar_bd()
                    cur = con.cursor()
                    cur.execute("""
                        INSERT INTO productos (
                            codigo_producto, codigo_barra, descripcion,
                            stock_actual, stock_minimo, stock_maximo, stock_reposicion,
                            costo, margen_venta
                        ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)
                    """, (
                        codigo_producto_p, codigo_barra_p.strip() or None, descripcion_p.upper(),
                        stock_actual_p, stock_minimo_p, stock_maximo_p, stock_reposicion_p,
                        costo_p, margen_p,
                    ))
                    con.commit()
                    st.session_state["_gen_producto"] = st.session_state.get("_gen_producto", 0) + 1
                    st.session_state["_flash_productos"] = f"✅ Producto **{codigo_producto_p}** registrado correctamente."
                    st.rerun(scope="fragment")
                except pg_errors.UniqueViolation:
                    st.error("❌ Ese Código de Producto ya existe.")
                except Exception as e:
                    st.error(f"❌ Error: {e}")
                finally:
                    if "cur" in locals(): cur.close()
                    if "con" in locals(): con.close()

    # --------------------------------------------------------
    # MÓDULO B: MODIFICAR PRODUCTO EXISTENTE
    # --------------------------------------------------------
    elif st.session_state["menu_productos"] == "modificar":
        st.subheader("Modificar Producto Existente")

        try:
            con = conectar_bd()
            cur = con.cursor()
            cur.execute("SELECT codigo_producto, descripcion, stock_actual FROM productos ORDER BY id DESC LIMIT 15")
            filas_prod = cur.fetchall()
            if filas_prod:
                st.markdown("📋 **Últimos Códigos de Producto registrados:**")
                st.dataframe(pd.DataFrame(filas_prod, columns=["Código", "Descripción", "Stock"]), use_container_width=True, hide_index=True)
        except: pass
        finally:
            if "cur" in locals(): cur.close()
            if "con" in locals(): con.close()

        st.write("---")
        busq_key_producto = f"busq_producto_mod_{st.session_state.get('_gen_producto', 0)}"
        busq_producto = st.text_input(
            "Ingresa el Código de Producto o Código de Barra a buscar:", key=busq_key_producto
        )

        if st.button("🔎 Buscar Producto", use_container_width=True):
            if not busq_producto.strip():
                st.warning("⚠️ Ingresa un código.")
            else:
                try:
                    con = conectar_bd()
                    cur = con.cursor()
                    cur.execute(
                        "SELECT * FROM productos WHERE codigo_producto = %s OR codigo_barra = %s",
                        (busq_producto.strip().upper(), busq_producto.strip()),
                    )
                    cols = [d[0] for d in cur.description]
                    fila = cur.fetchone()
                    if fila:
                        st.session_state["producto_mod"] = dict(zip(cols, fila))
                    else:
                        st.error("❌ No se encontró ningún producto con ese código.")
                except Exception as e:
                    st.error(f"❌ Error: {e}")
                finally:
                    if "cur" in locals(): cur.close()
                    if "con" in locals(): con.close()

        if "producto_mod" in st.session_state:
            p = st.session_state["producto_mod"]
            st.success(f"✅ Producto encontrado: **{p['codigo_producto']}** — {p.get('descripcion')}")

            cpm1, cpm2 = st.columns(2)
            with cpm1:
                codigo_barra_m = st.text_input(
                    "Código de Barra", value=p.get("codigo_barra") or "", key="codigo_barra_mod"
                )
            with cpm2:
                descripcion_m = st.text_input(
                    "Descripción", value=p.get("descripcion") or "", key="descripcion_mod"
                )

            st.markdown("---")
            st.markdown("### 📊 Stock")

            stock_actual_actual = int(p.get("stock_actual") or 0)
            stock_min_actual = int(p.get("stock_minimo") or 0)
            stock_max_actual = int(p.get("stock_maximo") or 0)
            stock_repo_actual = int(p.get("stock_reposicion") or 0)

            sm1, sm2, sm3, sm4 = st.columns(4)
            sm1.metric("Stock Actual", stock_actual_actual)
            sm2.metric("Stock Mínimo", stock_min_actual)
            sm3.metric("Stock Reposición", stock_repo_actual)
            sm4.metric("Stock Máximo", stock_max_actual)

            if stock_actual_actual <= stock_min_actual:
                st.error("🔴 Stock en el mínimo o bajo el mínimo — hay que reponer con urgencia.")
            elif stock_actual_actual <= stock_repo_actual:
                st.warning("🟡 Stock en el punto de reposición — es momento de empezar a comprar.")
            elif stock_max_actual and stock_actual_actual > stock_max_actual:
                st.info("🔵 Stock por sobre el máximo definido.")
            else:
                st.success("🟢 Stock en rango normal.")

            smov1, smov2 = st.columns(2)
            with smov1:
                cantidad_mov = st.number_input(
                    "Cantidad a mover", min_value=1, step=1, value=1, key="cantidad_mov_stock"
                )
            with smov2:
                nota_mov = st.text_input("Motivo (opcional)", key="nota_mov_stock")

            bmov1, bmov2 = st.columns(2)
            with bmov1:
                agregar_stock = st.button("➕ Agregar Stock", key="btn_agregar_stock", use_container_width=True)
            with bmov2:
                restar_stock = st.button("➖ Restar Stock", key="btn_restar_stock", use_container_width=True)

            if agregar_stock or restar_stock:
                delta = cantidad_mov if agregar_stock else -cantidad_mov
                nuevo_stock = stock_actual_actual + delta
                if nuevo_stock < 0:
                    st.error("❌ No puedes restar más stock del que hay disponible.")
                else:
                    try:
                        con = conectar_bd()
                        cur = con.cursor()
                        cur.execute("UPDATE productos SET stock_actual=%s WHERE id=%s", (nuevo_stock, p["id"]))
                        cur.execute(
                            "INSERT INTO movimientos_stock (producto_id, tipo, cantidad, nota) VALUES (%s,%s,%s,%s)",
                            (p["id"], "entrada" if agregar_stock else "salida", cantidad_mov, nota_mov or None),
                        )
                        con.commit()
                        st.session_state["producto_mod"]["stock_actual"] = nuevo_stock
                        st.session_state["_flash_stock"] = f"✅ Stock actualizado: {stock_actual_actual} → {nuevo_stock}"
                        st.rerun(scope="fragment")
                    except Exception as e:
                        st.error(f"❌ Error: {e}")
                    finally:
                        if "cur" in locals(): cur.close()
                        if "con" in locals(): con.close()

            if st.session_state.get("_flash_stock"):
                st.success(st.session_state.pop("_flash_stock"))

            st.markdown("**Ajustar umbrales de stock**")
            sem1, sem2, sem3 = st.columns(3)
            with sem1:
                stock_minimo_m = st.number_input(
                    "Stock Mínimo", min_value=0, step=1, value=stock_min_actual, key="stock_minimo_mod"
                )
            with sem2:
                stock_reposicion_m = st.number_input(
                    "Stock Reposición", min_value=0, step=1, value=stock_repo_actual, key="stock_reposicion_mod"
                )
            with sem3:
                stock_maximo_m = st.number_input(
                    "Stock Máximo", min_value=0, step=1, value=stock_max_actual, key="stock_maximo_mod"
                )

            st.markdown("---")
            st.markdown("### 💰 Precio de Venta")
            pm1, pm2, pm3, pm4 = st.columns(4)
            with pm1:
                costo_inicial_m = f"{int(p.get('costo') or 0):,}".replace(",", ".")
                txt_costo_m = st.text_input(
                    "Costo ($)", value=costo_inicial_m, key="costo_prod_mod",
                    on_change=formatear_monto_callback, args=("costo_prod_mod",),
                )
                costo_m = int(txt_costo_m.replace(".", "").replace("$", "").strip() or 0)
                st.caption(f"🔎 Sombra: **${costo_m:,.0f}**".replace(",", "."))
            with pm2:
                margen_m = st.number_input(
                    "Margen (%)", min_value=0.0, step=0.5, format="%.2f",
                    value=float(p.get("margen_venta") or 0), key="margen_prod_mod",
                )
            with pm3:
                neto_m = round(costo_m * (1 + margen_m / 100))
                st.metric("Neto Afecto a IVA", f"${neto_m:,.0f}".replace(",", "."))
            with pm4:
                precio_iva_m = round(neto_m * (1 + IVA_RATE))
                st.metric("Precio con IVA", f"${precio_iva_m:,.0f}".replace(",", "."))

            st.write("")
            actualizar_producto = st.button("💾 Actualizar Producto", use_container_width=True, type="primary")

            if actualizar_producto:
                try:
                    con = conectar_bd()
                    cur = con.cursor()
                    cur.execute("""
                        UPDATE productos SET
                            codigo_barra=%s, descripcion=%s,
                            stock_minimo=%s, stock_maximo=%s, stock_reposicion=%s,
                            costo=%s, margen_venta=%s
                        WHERE id=%s
                    """, (
                        codigo_barra_m.strip() or None, descripcion_m.upper(),
                        stock_minimo_m, stock_maximo_m, stock_reposicion_m,
                        costo_m, margen_m, p["id"],
                    ))
                    con.commit()
                    del st.session_state["producto_mod"]
                    st.session_state["_gen_producto"] = st.session_state.get("_gen_producto", 0) + 1
                    st.session_state["_flash_productos"] = "✅ ¡Producto actualizado exitosamente!"
                    st.rerun(scope="fragment")
                except Exception as e:
                    st.error(f"❌ Error: {e}")
                finally:
                    if "cur" in locals(): cur.close()
                    if "con" in locals(): con.close()

if "productos" in tabs:
    with tabs["productos"]:
        render_productos()

# ================================================================
# PESTAÑA 3 — GESTIÓN DE CLIENTES
# ================================================================
@st.fragment
def render_clientes():
    if "menu_clientes" not in st.session_state:
        st.session_state["menu_clientes"] = "crear"

    st.markdown("**¿Qué deseas hacer en Clientes?**")
    col_cli1, col_cli2 = st.columns(2)
    with col_cli1:
        if st.button("➕ Registrar nuevo cliente", key="btn_cli_crear", use_container_width=True):
            st.session_state["menu_clientes"] = "crear"
    with col_cli2:
        if st.button("✏️ Modificar cliente existente", key="btn_cli_mod", use_container_width=True):
            st.session_state["menu_clientes"] = "modificar"

    if st.session_state.get("_flash_clientes"):
        st.success(st.session_state.pop("_flash_clientes"))
    st.write("")

    # --------------------------------------------------------
    # MÓDULO A: REGISTRAR NUEVO CLIENTE
    # --------------------------------------------------------
    if st.session_state["menu_clientes"] == "crear":
        st.subheader("Registrar Nuevo Cliente")
        
        st.markdown("**Datos obligatorios**")
        c1, c2 = st.columns(2)
        with c1:
            rut_n = st.text_input("RUT * (Ej: 16328681-5 o 163286815)")
            razon_n = st.text_input("Razón Social *")
            fantasia_n = st.text_input("Nombre Fantasía")
            giro_n = st.text_input("Giro")
            direccion_n = st.text_input("Dirección")
            comuna_n = st.text_input("Comuna")
        with c2:
            ciudad_n = st.text_input("Ciudad")
            telefono_n = st.text_input("Teléfono")
            telefono2_n = st.text_input("Teléfono 2")
            email_n = st.text_input("E-Mail")
            web_n = st.text_input("Web")
            contacto_n = st.text_input("Contacto Comercial")

        st.markdown("**Datos comerciales**")
        observaciones_n = st.text_area("Observaciones")

        st.write("")
        # Botón de guardado libre de st.form
        guardar_cliente = st.button("💾 Guardar Cliente", use_container_width=True, type="primary")

        if guardar_cliente:
            if not rut_n or not razon_n:
                st.error("⚠️ El RUT y la Razón Social son obligatorios.")
            else:
                # Limpieza inteligente del RUT para validar
                rut_limpio = rut_n.replace(".","").replace("-","").strip()
                v_rut, r_rut = validar_rut(rut_limpio)
                
                if not v_rut:
                    st.error("❌ El RUT del Cliente es inválido.")
                else:
                    try:
                        # validar_rut ya devuelve el RUT formateado como "cuerpo-dv"
                        con = conectar_bd()
                        cur = con.cursor()
                        cur.execute("""
                            INSERT INTO clientes (
                                rut, razon_social, nombre_fantasia, giro, direccion, comuna, ciudad,
                                telefono, telefono2, email, web, contacto_comercial, observaciones
                            ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                        """, (
                            r_rut, razon_n.upper(), fantasia_n.upper() or None, giro_n.upper() or None,
                            direccion_n or None, comuna_n.upper() or None, ciudad_n.upper() or None,
                            telefono_n or None, telefono2_n or None, email_n or None, web_n or None, contacto_n or None,
                            observaciones_n or None
                        ))
                        con.commit()
                        st.success(f"✅ Cliente **{razon_n.upper()}** ({r_rut}) registrado correctamente.")
                    except pg_errors.UniqueViolation: 
                        st.error("❌ Ese RUT ya existe en el sistema.")
                    except Exception as e: 
                        st.error(f"❌ Error: {e}")
                    finally:
                        if "cur" in locals(): cur.close()
                        if "con" in locals(): con.close()

    # --------------------------------------------------------
    # MÓDULO B: MODIFICAR CLIENTE EXISTENTE
    # --------------------------------------------------------
    elif st.session_state["menu_clientes"] == "modificar":
        st.subheader("Modificar Cliente Existente")
        rut_key_cliente = f"rut_modificar_{st.session_state.get('_gen_cliente', 0)}"
        rut_mod = st.text_input("Ingresa el RUT a buscar (con o sin guion):", key=rut_key_cliente)

        if st.button("🔎 Buscar Cliente", use_container_width=True):
            if not rut_mod.strip(): 
                st.warning("⚠️ Ingresa un RUT válido.")
            else:
                # Transformamos la búsqueda para que calce exacto con la base de datos
                rut_limpio = rut_mod.replace(".", "").replace("-", "").strip()
                if len(rut_limpio) > 1:
                    rut_formateado = f"{rut_limpio[:-1]}-{rut_limpio[-1]}"
                else:
                    rut_formateado = rut_mod

                try:
                    con = conectar_bd()
                    cur = con.cursor()
                    cur.execute("SELECT * FROM clientes WHERE rut = %s", (rut_formateado,))
                    cols = [d[0] for d in cur.description]
                    fila = cur.fetchone()
                    
                    if fila:
                        st.session_state["cliente_mod"] = dict(zip(cols, fila))
                        st.success("✅ Cliente encontrado. Modifica los datos abajo:")
                    else: 
                        st.error(f"❌ No se encontró ningún cliente con el RUT {rut_formateado}.")
                except Exception as e: 
                    st.error(f"Error: {e}")
                finally:
                    if "cur" in locals(): cur.close()
                    if "con" in locals(): con.close()

        if "cliente_mod" in st.session_state:
            c = st.session_state["cliente_mod"]
            
            c1, c2 = st.columns(2)
            with c1:
                razon_m = st.text_input("Razón Social *", value=c.get("razon_social") or "")
                fantasia_m = st.text_input("Nombre Fantasía", value=c.get("nombre_fantasia") or "")
                giro_m = st.text_input("Giro", value=c.get("giro") or "")
                dir_m = st.text_input("Dirección", value=c.get("direccion") or "")
                comuna_m = st.text_input("Comuna", value=c.get("comuna") or "")
            with c2:
                ciudad_m = st.text_input("Ciudad", value=c.get("ciudad") or "")
                tel_m = st.text_input("Teléfono", value=c.get("telefono") or "")
                tel2_m = st.text_input("Teléfono 2", value=c.get("telefono2") or "")
                email_m = st.text_input("E-Mail", value=c.get("email") or "")
                obs_m = st.text_area("Observaciones", value=c.get("observaciones") or "")
            
            st.write("")
            actualizar = st.button("💾 Actualizar Cliente", use_container_width=True, type="primary")

            if actualizar:
                if not razon_m.strip():
                    st.error("⚠️ La Razón Social no puede estar vacía.")
                else:
                    try:
                        con = conectar_bd()
                        cur = con.cursor()
                        cur.execute("""
                            UPDATE clientes SET
                                razon_social=%s, nombre_fantasia=%s, giro=%s, direccion=%s, comuna=%s, ciudad=%s,
                                telefono=%s, telefono2=%s, email=%s, observaciones=%s
                            WHERE rut=%s
                        """, (
                            razon_m.upper(), fantasia_m.upper() or None, giro_m.upper() or None,
                            dir_m or None, comuna_m.upper() or None, ciudad_m.upper() or None,
                            tel_m or None, tel2_m or None, email_m or None, obs_m or None, c["rut"]
                        ))
                        con.commit()
                        del st.session_state["cliente_mod"]
                        st.session_state["_gen_cliente"] = st.session_state.get("_gen_cliente", 0) + 1
                        st.session_state["_flash_clientes"] = "✅ ¡Cliente actualizado exitosamente!"
                        st.rerun(scope="fragment")
                    except Exception as e:
                        st.error(f"❌ Error: {e}")
                    finally:
                        if "cur" in locals(): cur.close()
                        if "con" in locals(): con.close()

if "clientes" in tabs:
    with tabs["clientes"]:
        render_clientes()

# ================================================================
# PESTAÑA 4 — GESTIÓN DE CRÉDITOS
# ================================================================
@st.fragment
def render_creditos():
    if "menu_creditos" not in st.session_state:
        st.session_state["menu_creditos"] = "crear"

    st.markdown("**¿Qué deseas hacer en Créditos?**")
    col_cre1, col_cre2 = st.columns(2)
    with col_cre1:
        if st.button("➕ Registrar nuevo crédito", key="btn_cre_crear", use_container_width=True):
            st.session_state["menu_creditos"] = "crear"
    with col_cre2:
        if st.button("✏️ Modificar crédito existente", key="btn_cre_mod", use_container_width=True):
            st.session_state["menu_creditos"] = "modificar"

    if st.session_state.get("_flash_creditos_menu"):
        st.success(st.session_state.pop("_flash_creditos_menu"))
    st.write("")

    # --------------------------------------------------------
    # MÓDULO A: REGISTRAR NUEVO CRÉDITO
    # --------------------------------------------------------
    if st.session_state["menu_creditos"] == "crear":
        st.subheader("Registrar Nuevo Crédito")

        st.markdown("**🚗 Vehículo y Cliente**")
        cr1, cr2 = st.columns(2)
        with cr1:
            patente_cr = st.text_input("Patente *", key="patente_cre_crear")
            rut_cliente_cr = st.text_input("RUT Cliente * (Acepta con o sin guion)", key="rut_cre_crear")
        with cr2:
            rut_limpio_lookup_cr = rut_cliente_cr.replace(".", "").replace("-", "").strip()
            cliente_encontrado_cr = None
            # Key dinámica: cambia con el RUT para forzar que el widget refleje el valor actual
            # (Streamlit solo aplica `value=` la primera vez que se crea un widget con una key fija).
            nombre_key_cr = f"nombre_cre_crear_auto_{rut_limpio_lookup_cr or 'vacio'}"
            if not rut_limpio_lookup_cr:
                st.text_input(
                    "Nombre de Cliente", value="", disabled=True, key=nombre_key_cr,
                    placeholder="Ingresa el RUT para autocompletar",
                )
            else:
                v_rut_lookup_cr, r_rut_lookup_cr = validar_rut(rut_limpio_lookup_cr)
                if not v_rut_lookup_cr:
                    st.text_input("Nombre de Cliente", value="", disabled=True, key=nombre_key_cr)
                    st.caption("⚠️ El RUT ingresado no es válido.")
                else:
                    try:
                        con = conectar_bd()
                        cur = con.cursor()
                        cur.execute("SELECT razon_social FROM clientes WHERE rut=%s", (r_rut_lookup_cr,))
                        fila = cur.fetchone()
                        cliente_encontrado_cr = fila[0] if fila else None
                    finally:
                        if "cur" in locals(): cur.close()
                        if "con" in locals(): con.close()

                    if cliente_encontrado_cr:
                        st.text_input(
                            "Nombre de Cliente", value=cliente_encontrado_cr, disabled=True, key=nombre_key_cr
                        )
                    else:
                        st.text_input("Nombre de Cliente", value="", disabled=True, key=nombre_key_cr)
                        st.caption("❌ Cliente no registrado. Regístralo primero en la pestaña Clientes.")

        st.markdown("---")
        st.markdown("### 💰 Condiciones del Crédito")
        st.markdown("_Puedes ingresar los montos con o sin puntos (Ej: 10.000.000 o 10000000)_")
        ce1, ce2, ce3 = st.columns(3)
        with ce1:
            txt_precio_cr = st.text_input("Precio Venta ($) *", value="0", key="txt_precio_crear")
            precio_venta_cr = int(txt_precio_cr.replace(".", "").replace("$", "").strip() or 0)
            st.caption(f"🔎 Sombra: **${precio_venta_cr:,.0f}**".replace(",", "."))
        with ce2:
            txt_pie_cr = st.text_input("Abono / Pie ($)", value="0", key="txt_pie_crear")
            pie_cr = int(txt_pie_cr.replace(".", "").replace("$", "").strip() or 0)
            st.caption(f"🔎 Sombra: **${pie_cr:,.0f}**".replace(",", "."))
        with ce3:
            saldo_financiamiento_cr = precio_venta_cr - pie_cr
            st.metric("Saldo Financiamiento", f"${saldo_financiamiento_cr:,.0f}".replace(",", "."))

        cc1, cc2 = st.columns(2)
        with cc1:
            cantidad_cuotas_cr = st.number_input("Cantidad de Cuotas", min_value=0, step=1, key="cant_cuotas_crear")
        with cc2:
            valor_cuota_cr = round(saldo_financiamiento_cr / cantidad_cuotas_cr) if cantidad_cuotas_cr else 0
            st.metric("Valor Cuota", f"${valor_cuota_cr:,.0f}".replace(",", "."))

        st.markdown("---")
        st.markdown("### 🧾 Comisión")
        cm1, cm2 = st.columns(2)
        with cm1:
            txt_comision_cr = st.text_input("Monto Comisión ($)", value="0", key="txt_comision_crear")
            monto_comision_cr = int(txt_comision_cr.replace(".", "").replace("$", "").strip() or 0)
            st.caption(f"🔎 Sombra: **${monto_comision_cr:,.0f}**".replace(",", "."))
        with cm2:
            estado_comision_cr = st.radio(
                "Estado de la Comisión", ["Pendiente", "Pagada"], horizontal=True, key="estado_comision_crear"
            )

        if estado_comision_cr == "Pagada":
            ccp1, ccp2 = st.columns(2)
            with ccp1:
                forma_pago_comision_cr = st.selectbox(
                    "Forma de Pago de la Comisión", FORMAS_PAGO_DISPONIBLES, key="forma_pago_comision_crear"
                )
            with ccp2:
                documentos_comision_cr = st.file_uploader(
                    "Documento(s) de Pago de Comisión", key="documento_comision_crear",
                    accept_multiple_files=True,
                )
                st.caption("Puedes seleccionar varios archivos a la vez (ej: varias transferencias).")
        else:
            forma_pago_comision_cr = None
            documentos_comision_cr = []

        st.markdown("---")
        st.markdown("### 💳 Forma de Pago del Pie / Abono")
        forma_pago_cr = st.multiselect(
            "Selecciona una o más formas de pago", FORMAS_PAGO_DISPONIBLES, key="forma_pago_crear"
        )

        st.markdown("---")
        st.markdown("### 📎 Documento(s) de Respaldo General")
        documentos_cr = st.file_uploader(
            "Adjuntar documento(s) (PDF, imagen, Word, etc.)", key="documento_crear",
            accept_multiple_files=True,
        )
        st.caption("Puedes seleccionar varios archivos a la vez.")

        st.write("")
        guardar_credito = st.button("💾 Guardar Crédito", use_container_width=True, type="primary")

        if guardar_credito:
            rut_limpio_cr = rut_cliente_cr.replace(".", "").replace("-", "").strip()

            if not patente_cr.strip() or not rut_limpio_cr:
                st.error("⚠️ La Patente y el RUT del Cliente son obligatorios.")
            else:
                v_rut_cr, r_rut_cr = validar_rut(rut_limpio_cr)
                if not v_rut_cr:
                    st.error("❌ El RUT del Cliente ingresado no es válido.")
                elif not cliente_encontrado_cr:
                    st.error(f"❌ No existe un cliente con RUT {r_rut_cr}. Regístralo primero en la pestaña Clientes.")
                else:
                    patente_limpia_cr = patente_cr.upper().replace("-", "").replace(" ", "").strip()
                    try:
                        con = conectar_bd()
                        cur = con.cursor()

                        cur.execute("SELECT 1 FROM vehiculos WHERE patente=%s", (patente_limpia_cr,))
                        if not cur.fetchone():
                            st.error(f"❌ No existe un vehículo registrado con la patente {patente_limpia_cr}. Regístralo primero en la pestaña Vehículos.")
                        else:
                            cur.execute("""
                                INSERT INTO creditos (
                                    patente, nombre_cliente, rut_cliente, precio_venta, pie, saldo_financiamiento,
                                    cantidad_cuotas, cuotas_pendientes, valor_cuota, monto_comision, comision_pagada, forma_pago
                                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                                RETURNING id
                            """, (
                                patente_limpia_cr, cliente_encontrado_cr, r_rut_cr,
                                precio_venta_cr, pie_cr, saldo_financiamiento_cr,
                                cantidad_cuotas_cr, cantidad_cuotas_cr, valor_cuota_cr, monto_comision_cr,
                                estado_comision_cr == "Pagada", ", ".join(forma_pago_cr) or None,
                            ))
                            nuevo_credito_id = cur.fetchone()[0]

                            for doc in (documentos_cr or []):
                                cur.execute("""
                                    INSERT INTO documentos_credito (credito_id, categoria, documento_nombre, documento_tipo, documento_datos)
                                    VALUES (%s,'general',%s,%s,%s)
                                """, (nuevo_credito_id, doc.name, doc.type, psycopg2.Binary(doc.getvalue())))

                            if estado_comision_cr == "Pagada":
                                cur.execute("""
                                    INSERT INTO pagos_credito (credito_id, tipo_pago, monto, forma_pago)
                                    VALUES (%s,'comision',%s,%s)
                                """, (nuevo_credito_id, monto_comision_cr, forma_pago_comision_cr))

                                for doc in (documentos_comision_cr or []):
                                    cur.execute("""
                                        INSERT INTO documentos_credito (credito_id, categoria, documento_nombre, documento_tipo, documento_datos)
                                        VALUES (%s,'comision',%s,%s,%s)
                                    """, (nuevo_credito_id, doc.name, doc.type, psycopg2.Binary(doc.getvalue())))

                            con.commit()
                            st.success(f"✅ Crédito del vehículo **{patente_limpia_cr}** registrado correctamente.")
                    except Exception as e:
                        st.error(f"❌ Error: {e}")
                    finally:
                        if "cur" in locals(): cur.close()
                        if "con" in locals(): con.close()

    # --------------------------------------------------------
    # MÓDULO B: MODIFICAR CRÉDITO EXISTENTE
    # --------------------------------------------------------
    elif st.session_state["menu_creditos"] == "modificar":
        st.subheader("Modificar Crédito Existente")
        patente_key_credito = f"patente_mod_creditos_{st.session_state.get('_gen_credito', 0)}"
        patente_mod_cr = st.text_input("Ingresa la Patente del crédito a buscar:", key=patente_key_credito)

        if st.button("🔎 Buscar Crédito", use_container_width=True):
            if not patente_mod_cr.strip():
                st.warning("⚠️ Ingresa una patente.")
            else:
                patente_busq = patente_mod_cr.upper().replace("-", "").replace(" ", "").strip()
                try:
                    con = conectar_bd()
                    cur = con.cursor()
                    cur.execute(
                        "SELECT * FROM creditos WHERE patente = %s ORDER BY creado_en DESC LIMIT 1",
                        (patente_busq,),
                    )
                    cols = [d[0] for d in cur.description]
                    fila = cur.fetchone()
                    if fila:
                        st.session_state["credito_mod"] = dict(zip(cols, fila))
                        st.session_state.pop("mostrar_pago_comision", None)
                        st.session_state.pop("mostrar_pago_cuota", None)
                    else:
                        st.error(f"❌ No se encontró ningún crédito para la patente {patente_busq}.")
                except Exception as e:
                    st.error(f"❌ Error: {e}")
                finally:
                    if "cur" in locals(): cur.close()
                    if "con" in locals(): con.close()

        if "credito_mod" in st.session_state:
            cr = st.session_state["credito_mod"]
            st.success(f"✅ Crédito encontrado ({cr['patente']}). Modifica los datos abajo:")

            # Historial de pagos y documentos de este crédito
            try:
                con = conectar_bd()
                cur = con.cursor()
                cur.execute("SELECT * FROM pagos_credito WHERE credito_id=%s ORDER BY fecha_pago", (cr["id"],))
                pago_cols = [d[0] for d in cur.description]
                pagos = [dict(zip(pago_cols, fila)) for fila in cur.fetchall()]

                cur.execute("SELECT * FROM documentos_credito WHERE credito_id=%s ORDER BY creado_en", (cr["id"],))
                doc_cols = [d[0] for d in cur.description]
                documentos = [dict(zip(doc_cols, fila)) for fila in cur.fetchall()]
            except Exception as e:
                st.error(f"❌ Error al cargar el historial de pagos: {e}")
                pagos = []
                documentos = []
            finally:
                if "cur" in locals(): cur.close()
                if "con" in locals(): con.close()

            pagos_cuota = [p for p in pagos if p["tipo_pago"] == "cuota"]
            pago_comision = next((p for p in pagos if p["tipo_pago"] == "comision"), None)
            cuotas_pagadas = len(pagos_cuota)
            monto_pagado_cuotas = sum(int(p["monto"] or 0) for p in pagos_cuota)
            documentos_comision = [d for d in documentos if d["categoria"] == "comision"]
            documentos_general = [d for d in documentos if d["categoria"] == "general"]

            crm1, crm2 = st.columns(2)
            with crm1:
                nombre_cliente_m = st.text_input("Nombre de Cliente", value=cr.get("nombre_cliente") or "", key="nombre_cre_mod")
            with crm2:
                rut_cliente_m = st.text_input("RUT Cliente", value=cr.get("rut_cliente") or "", key="rut_cre_mod")

            st.markdown("---")
            st.markdown("### 💰 Condiciones del Crédito")
            cem1, cem2, cem3 = st.columns(3)
            with cem1:
                precio_inicial_m = f"{int(cr.get('precio_venta') or 0):,}".replace(",", ".")
                txt_precio_m = st.text_input("Precio Venta ($)", value=precio_inicial_m, key="txt_precio_mod")
                precio_venta_m = int(txt_precio_m.replace(".", "").replace("$", "").strip() or 0)
                st.caption(f"🔎 Sombra: **${precio_venta_m:,.0f}**".replace(",", "."))
            with cem2:
                pie_inicial_m = f"{int(cr.get('pie') or 0):,}".replace(",", ".")
                txt_pie_m = st.text_input("Abono / Pie ($)", value=pie_inicial_m, key="txt_pie_mod")
                pie_m = int(txt_pie_m.replace(".", "").replace("$", "").strip() or 0)
                st.caption(f"🔎 Sombra: **${pie_m:,.0f}**".replace(",", "."))
            with cem3:
                saldo_financiamiento_m = precio_venta_m - pie_m
                st.metric("Saldo Financiamiento", f"${saldo_financiamiento_m:,.0f}".replace(",", "."))

            ccm1, ccm2 = st.columns(2)
            with ccm1:
                cantidad_cuotas_m = st.number_input(
                    "Cantidad de Cuotas", value=int(cr.get("cantidad_cuotas") or 0), min_value=0, step=1, key="cant_cuotas_mod"
                )
            with ccm2:
                valor_cuota_m = round(saldo_financiamiento_m / cantidad_cuotas_m) if cantidad_cuotas_m else 0
                st.metric("Valor Cuota", f"${valor_cuota_m:,.0f}".replace(",", "."))

            st.markdown("---")
            st.markdown("### 🧾 Comisión")
            comision_inicial_m = f"{int(cr.get('monto_comision') or 0):,}".replace(",", ".")
            txt_comision_m = st.text_input("Monto Comisión ($)", value=comision_inicial_m, key="txt_comision_mod")
            monto_comision_m = int(txt_comision_m.replace(".", "").replace("$", "").strip() or 0)
            st.caption(f"🔎 Sombra: **${monto_comision_m:,.0f}**".replace(",", "."))

            if cr.get("comision_pagada"):
                st.success("✅ Comisión Pagada")
                if pago_comision:
                    st.caption(
                        f"Pagada el {pago_comision['fecha_pago'].strftime('%d-%m-%Y')} · "
                        f"Forma de pago: {pago_comision.get('forma_pago') or 'No indicada'}"
                    )
                if documentos_comision:
                    st.markdown("**Comprobantes de comisión**")
                    for doc in documentos_comision:
                        dc1, dc2 = st.columns([3, 1])
                        dc1.write(f"📄 {doc.get('documento_nombre')}")
                        with dc2:
                            st.download_button(
                                "📥 Descargar",
                                data=bytes(doc["documento_datos"]),
                                file_name=doc.get("documento_nombre") or "comprobante_comision",
                                mime=doc.get("documento_tipo") or "application/octet-stream",
                                key=f"descargar_doc_comision_{doc['id']}",
                            )
                else:
                    st.caption("Sin comprobantes adjuntos.")

                nuevos_docs_comision = st.file_uploader(
                    "➕ Agregar más comprobantes de comisión", key=f"agregar_doc_comision_{len(documentos_comision)}",
                    accept_multiple_files=True,
                )
                if st.button("➕ Adjuntar Documento(s) de Comisión", key="btn_adjuntar_doc_comision"):
                    if not nuevos_docs_comision:
                        st.warning("⚠️ Selecciona al menos un archivo primero.")
                    else:
                        try:
                            con = conectar_bd()
                            cur = con.cursor()
                            for doc in nuevos_docs_comision:
                                cur.execute("""
                                    INSERT INTO documentos_credito (credito_id, categoria, documento_nombre, documento_tipo, documento_datos)
                                    VALUES (%s,'comision',%s,%s,%s)
                                """, (cr["id"], doc.name, doc.type, psycopg2.Binary(doc.getvalue())))
                            con.commit()
                            st.session_state["_flash_doc_comision"] = "✅ Documento(s) adjuntado(s)."
                            st.rerun(scope="fragment")
                        except Exception as e:
                            st.error(f"❌ Error: {e}")
                        finally:
                            if "cur" in locals(): cur.close()
                            if "con" in locals(): con.close()

                if st.session_state.get("_flash_doc_comision"):
                    st.success(st.session_state.pop("_flash_doc_comision"))
            else:
                st.warning("🕒 Comisión Pendiente")
                if st.button("💰 ¿Desea pagar la comisión?", key="btn_mostrar_pago_comision"):
                    st.session_state["mostrar_pago_comision"] = True

                if st.session_state.get("mostrar_pago_comision"):
                    forma_pago_comision_m = st.selectbox(
                        "Forma de Pago de la Comisión", FORMAS_PAGO_DISPONIBLES, key="forma_pago_comision_mod"
                    )
                    documentos_comision_m = st.file_uploader(
                        "Documento(s) de Pago de Comisión", key="documento_comision_mod",
                        accept_multiple_files=True,
                    )
                    st.caption("Puedes seleccionar varios archivos a la vez (ej: varias transferencias).")
                    if st.button("✅ Confirmar Pago de Comisión", key="btn_confirmar_pago_comision", type="primary"):
                        try:
                            con = conectar_bd()
                            cur = con.cursor()
                            cur.execute("""
                                INSERT INTO pagos_credito (credito_id, tipo_pago, monto, forma_pago)
                                VALUES (%s,'comision',%s,%s)
                            """, (cr["id"], monto_comision_m, forma_pago_comision_m))
                            for doc in (documentos_comision_m or []):
                                cur.execute("""
                                    INSERT INTO documentos_credito (credito_id, categoria, documento_nombre, documento_tipo, documento_datos)
                                    VALUES (%s,'comision',%s,%s,%s)
                                """, (cr["id"], doc.name, doc.type, psycopg2.Binary(doc.getvalue())))
                            cur.execute(
                                "UPDATE creditos SET comision_pagada=TRUE, monto_comision=%s WHERE id=%s",
                                (monto_comision_m, cr["id"]),
                            )
                            con.commit()
                            st.session_state.pop("mostrar_pago_comision", None)
                            st.session_state["credito_mod"]["comision_pagada"] = True
                            st.session_state["credito_mod"]["monto_comision"] = monto_comision_m
                            st.session_state["_flash_pago_comision"] = "✅ Pago de comisión registrado."
                            st.rerun(scope="fragment")
                        except Exception as e:
                            st.error(f"❌ Error: {e}")
                        finally:
                            if "cur" in locals(): cur.close()
                            if "con" in locals(): con.close()

            if st.session_state.get("_flash_pago_comision"):
                st.success(st.session_state.pop("_flash_pago_comision"))

            st.markdown("---")
            st.markdown("### 💳 Forma de Pago del Pie / Abono")
            forma_pago_previa = [f.strip() for f in (cr.get("forma_pago") or "").split(",") if f.strip()]
            forma_pago_m = st.multiselect(
                "Selecciona una o más formas de pago", FORMAS_PAGO_DISPONIBLES,
                default=[f for f in forma_pago_previa if f in FORMAS_PAGO_DISPONIBLES],
                key="forma_pago_mod",
            )

            st.markdown("---")
            st.markdown("### 📎 Documento(s) de Respaldo General")
            if documentos_general:
                for doc in documentos_general:
                    dg1, dg2 = st.columns([3, 1])
                    dg1.write(f"📄 {doc.get('documento_nombre')}")
                    with dg2:
                        st.download_button(
                            "📥 Descargar",
                            data=bytes(doc["documento_datos"]),
                            file_name=doc.get("documento_nombre") or "documento",
                            mime=doc.get("documento_tipo") or "application/octet-stream",
                            key=f"descargar_doc_general_{doc['id']}",
                        )
            else:
                st.caption("Sin documentos adjuntos.")

            nuevos_docs_general = st.file_uploader(
                "➕ Agregar documento(s) de respaldo", key=f"agregar_doc_general_{len(documentos_general)}",
                accept_multiple_files=True,
            )
            if st.button("➕ Adjuntar Documento(s)", key="btn_adjuntar_doc_general"):
                if not nuevos_docs_general:
                    st.warning("⚠️ Selecciona al menos un archivo primero.")
                else:
                    try:
                        con = conectar_bd()
                        cur = con.cursor()
                        for doc in nuevos_docs_general:
                            cur.execute("""
                                INSERT INTO documentos_credito (credito_id, categoria, documento_nombre, documento_tipo, documento_datos)
                                VALUES (%s,'general',%s,%s,%s)
                            """, (cr["id"], doc.name, doc.type, psycopg2.Binary(doc.getvalue())))
                        con.commit()
                        st.session_state["_flash_doc_general"] = "✅ Documento(s) adjuntado(s)."
                        st.rerun(scope="fragment")
                    except Exception as e:
                        st.error(f"❌ Error: {e}")
                    finally:
                        if "cur" in locals(): cur.close()
                        if "con" in locals(): con.close()

            if st.session_state.get("_flash_doc_general"):
                st.success(st.session_state.pop("_flash_doc_general"))

            st.write("")
            actualizar_credito = st.button("💾 Actualizar Crédito", use_container_width=True, type="primary")

            if actualizar_credito:
                cuotas_pendientes_m = max(cantidad_cuotas_m - cuotas_pagadas, 0)
                try:
                    con = conectar_bd()
                    cur = con.cursor()
                    cur.execute("""
                        UPDATE creditos SET
                            nombre_cliente=%s, rut_cliente=%s, precio_venta=%s, pie=%s, saldo_financiamiento=%s,
                            cantidad_cuotas=%s, cuotas_pendientes=%s, valor_cuota=%s, monto_comision=%s, forma_pago=%s
                        WHERE id=%s
                    """, (
                        nombre_cliente_m.upper() or None, rut_cliente_m or None, precio_venta_m, pie_m, saldo_financiamiento_m,
                        cantidad_cuotas_m, cuotas_pendientes_m, valor_cuota_m, monto_comision_m,
                        ", ".join(forma_pago_m) or None,
                        cr["id"],
                    ))

                    con.commit()
                    del st.session_state["credito_mod"]
                    st.session_state["_gen_credito"] = st.session_state.get("_gen_credito", 0) + 1
                    st.session_state["_flash_creditos_menu"] = "✅ ¡Crédito actualizado exitosamente!"
                    st.rerun(scope="fragment")
                except Exception as e:
                    st.error(f"❌ Error: {e}")
                finally:
                    if "cur" in locals(): cur.close()
                    if "con" in locals(): con.close()

            st.markdown("---")
            st.markdown("### 📅 Pagos de Cuotas")
            saldo_pendiente_m = saldo_financiamiento_m - monto_pagado_cuotas
            pm1, pm2, pm3, pm4 = st.columns(4)
            pm1.metric("Cuotas Pagadas", f"{cuotas_pagadas} / {int(cr.get('cantidad_cuotas') or 0)}")
            pm2.metric("Valor Cuota", f"${valor_cuota_m:,.0f}".replace(",", "."))
            pm3.metric("Monto Pagado", f"${monto_pagado_cuotas:,.0f}".replace(",", "."))
            pm4.metric("Saldo Pendiente", f"${saldo_pendiente_m:,.0f}".replace(",", "."))

            if pagos_cuota:
                st.markdown("**Historial de pagos**")
                for i, p in enumerate(pagos_cuota, start=1):
                    hp1, hp2, hp3, hp4 = st.columns([1, 2, 2, 2])
                    hp1.write(f"Cuota {i}")
                    hp2.write(p["fecha_pago"].strftime("%d-%m-%Y"))
                    hp3.write(f"${int(p['monto'] or 0):,.0f}".replace(",", ".") + f" · {p.get('forma_pago') or ''}")
                    with hp4:
                        if p.get("documento_datos"):
                            st.download_button(
                                "📥 Comprobante",
                                data=bytes(p["documento_datos"]),
                                file_name=p.get("documento_nombre") or f"comprobante_cuota_{i}",
                                mime=p.get("documento_tipo") or "application/octet-stream",
                                key=f"descargar_cuota_{p['id']}",
                            )
                        else:
                            st.caption("Sin documento")

            if st.button("💵 ¿Desea pagar una cuota?", key="btn_mostrar_pago_cuota"):
                st.session_state["mostrar_pago_cuota"] = True

            if st.session_state.get("mostrar_pago_cuota"):
                st.markdown("**Registrar Pago de Cuota**")
                pcu1, pcu2 = st.columns(2)
                with pcu1:
                    txt_monto_cuota = st.text_input(
                        "Monto Pagado ($)", value=str(valor_cuota_m), key="txt_monto_cuota_pago"
                    )
                    monto_cuota_pago = int(txt_monto_cuota.replace(".", "").replace("$", "").strip() or 0)
                with pcu2:
                    forma_pago_cuota = st.selectbox(
                        "Forma de Pago", FORMAS_PAGO_DISPONIBLES, key="forma_pago_cuota_pago"
                    )
                documento_cuota = st.file_uploader(
                    "Documento de respaldo de esta cuota", key=f"documento_cuota_pago_{len(pagos_cuota)}"
                )
                if st.button("✅ Confirmar Pago de Cuota", key="btn_confirmar_pago_cuota", type="primary"):
                    try:
                        con = conectar_bd()
                        cur = con.cursor()
                        doc_cu_nombre = doc_cu_tipo = doc_cu_datos = None
                        if documento_cuota is not None:
                            doc_cu_nombre = documento_cuota.name
                            doc_cu_tipo = documento_cuota.type
                            doc_cu_datos = psycopg2.Binary(documento_cuota.getvalue())
                        cur.execute("""
                            INSERT INTO pagos_credito (
                                credito_id, tipo_pago, monto, forma_pago, documento_nombre, documento_tipo, documento_datos
                            ) VALUES (%s,'cuota',%s,%s,%s,%s,%s)
                        """, (
                            cr["id"], monto_cuota_pago, forma_pago_cuota,
                            doc_cu_nombre, doc_cu_tipo, doc_cu_datos,
                        ))
                        nuevas_cuotas_pagadas = cuotas_pagadas + 1
                        nuevas_cuotas_pendientes = max(int(cr.get("cantidad_cuotas") or 0) - nuevas_cuotas_pagadas, 0)
                        cur.execute(
                            "UPDATE creditos SET cuotas_pendientes=%s WHERE id=%s",
                            (nuevas_cuotas_pendientes, cr["id"]),
                        )
                        con.commit()
                        st.session_state["credito_mod"]["cuotas_pendientes"] = nuevas_cuotas_pendientes
                        st.session_state.pop("mostrar_pago_cuota", None)
                        st.session_state["_flash_pago_cuota"] = "✅ Pago de cuota registrado."
                        st.rerun(scope="fragment")
                    except Exception as e:
                        st.error(f"❌ Error: {e}")
                    finally:
                        if "cur" in locals(): cur.close()
                        if "con" in locals(): con.close()

            if st.session_state.get("_flash_pago_cuota"):
                st.success(st.session_state.pop("_flash_pago_cuota"))

if "creditos" in tabs:
    with tabs["creditos"]:
        render_creditos()

# ================================================================
# PESTAÑA 5 — GESTIÓN DE VEHÍCULOS
# ================================================================
@st.fragment
def render_vehiculos():
    if "menu_vehiculos" not in st.session_state:
        st.session_state["menu_vehiculos"] = "crear"

    st.markdown("**¿Qué deseas hacer en Vehículos?**")
    col_veh1, col_veh2 = st.columns(2)
    with col_veh1:
        if st.button("➕ Registrar nuevo vehículo", use_container_width=True):
            st.session_state["menu_vehiculos"] = "crear"
    with col_veh2:
        if st.button("✏️ Modificar vehículo existente", use_container_width=True):
            st.session_state["menu_vehiculos"] = "modificar"

    if st.session_state.get("_flash_vehiculos"):
        st.success(st.session_state.pop("_flash_vehiculos"))
    st.write("")

    # --------------------------------------------------------
    # MÓDULO A: REGISTRAR NUEVO VEHÍCULO
    # --------------------------------------------------------
    if st.session_state["menu_vehiculos"] == "crear":
        st.subheader("Registrar Nuevo Vehículo")
        st.info("💡 Escribe los datos libremente. Los valores económicos se sumarán automáticamente en tiempo real.")
        
        st.markdown("**🚗 Datos Técnicos del Vehículo**")
        v1, v2 = st.columns(2)
        with v1:
            patente_v = st.text_input("Patente *")
            marca_v = st.text_input("Marca")
            modelo_v = st.text_input("Modelo")
            año_v = st.number_input("Año", min_value=1900, max_value=2100, step=1, value=2026)
            km_v = st.number_input("Kilometraje", min_value=0, step=1000)
        with v2:
            chasis_v = st.text_input("Nro de Chasis")
            motor_v = st.text_input("Nro de Motor")
            puertas_v = st.number_input("Nro de Puertas", min_value=0, max_value=10, step=1, value=4)
            color_v = st.text_input("Color")

        st.markdown("**📍 Ubicación y Observaciones**")
        u1, u2 = st.columns(2)
        with u1:
            ubicacion_v = st.text_input("Ubicación", key="ubicacion_crear")
        with u2:
            observaciones_v = st.text_area("Observaciones", height=100, key="observaciones_crear")

        st.markdown("**🔖 Estado del Vehículo**")
        estado_reserva_v = st.radio(
            "Estado", ["Disponible", "Reservado"], horizontal=True, key="estado_reserva_crear"
        )
        if estado_reserva_v == "Reservado":
            r1, r2, r3 = st.columns(3)
            with r1:
                dias_reserva_v = st.number_input("Días de Reserva", min_value=0, step=1, key="dias_reserva_crear")
            with r2:
                monto_reserva_v = st.number_input("Monto de la Reserva ($)", min_value=0, step=1000, key="monto_reserva_crear")
            with r3:
                vendedora_v = st.text_input("Vendedora que atendió", key="vendedora_crear")
        else:
            dias_reserva_v = None
            monto_reserva_v = None
            vendedora_v = ""

        st.markdown("---")

        # --- SECCIÓN ECONÓMICA 100% EN VIVO ---
        st.markdown("### 💰 Detalles Económicos y Financieros")
        st.markdown("_Puedes ingresar los montos con o sin puntos (Ej: 10.000.000 o 10000000)_")
        col_eco1, col_eco2, col_eco3 = st.columns(3)
        
        with col_eco1:
            txt_valor = st.text_input(
                "Valor Vehículo ($) *", value="0", key="txt_val_crear",
                on_change=formatear_monto_callback, args=("txt_val_crear",),
            )
            valor_v = int(txt_valor.replace(".", "").replace("$", "").strip() or 0)
            st.caption(f"🔎 Sombra: **${valor_v:,.0f}**".replace(",", "."))

        with col_eco2:
            txt_comision = st.text_input(
                "Comisión ($) *", value="0", key="txt_com_crear",
                on_change=formatear_monto_callback, args=("txt_com_crear",),
            )
            comision_v = int(txt_comision.replace(".", "").replace("$", "").strip() or 0)
            st.caption(f"🔎 Sombra: **${comision_v:,.0f}**".replace(",", "."))

        with col_eco3:
            txt_transferencia = st.text_input(
                "Transferencia ($) *", value="0", key="txt_tra_crear",
                on_change=formatear_monto_callback, args=("txt_tra_crear",),
            )
            transferencia_v = int(txt_transferencia.replace(".", "").replace("$", "").strip() or 0)
            st.caption(f"🔎 Sombra: **${transferencia_v:,.0f}**".replace(",", "."))
        
        valor_final_v = valor_v + comision_v + transferencia_v
        st.success(f"💵 **Valor Final Calculado (En Tiempo Real):** ${valor_final_v:,.0f}".replace(",", "."))
        
        st.markdown("---")

        st.markdown("**👤 Datos del Cliente (quien trae el vehículo)**")
        rut_cli_v = st.text_input("RUT Cliente * (Acepta con o sin guion)")
        st.markdown("<p style='color: gray; font-size: 12px; margin-top: -10px;'>El cliente debe estar creado previamente en la pestaña Clientes.</p>", unsafe_allow_html=True)

        st.markdown("**📋 Datos del Dueño del Vehículo** *(si es distinto al cliente)*")
        e1, e2 = st.columns(2)
        with e1:
            nombre_dueno_v = st.text_input("Nombre Dueño Vehículo")
            rut_dueno_v = st.text_input("RUT Dueño Vehículo (Acepta con o sin guion)")
            dir_dueno_v = st.text_input("Dirección Dueño")
        with e2:
            tel_dueno_v = st.text_input("Teléfono Dueño")
            email_dueno_v = st.text_input("E-Mail Dueño")

        st.write("")
        if st.button("💾 Guardar Registro de Vehículo Completo", use_container_width=True, type="primary"):
            if not patente_v or not rut_cli_v:
                st.error("❌ Error: La Patente y el RUT del Cliente son obligatorios.")
            else:
                rut_cli_limpio = rut_cli_v.replace(".","").replace("-","").strip()
                rut_due_limpio = rut_dueno_v.replace(".","").replace("-","").strip()
                
                v_cli, r_cli = validar_rut(rut_cli_limpio)
                v_due, r_due = validar_rut(rut_due_limpio) if rut_due_limpio else (True, None)

                if not v_cli: st.error("❌ El RUT del Cliente ingresado no es válido.")
                elif rut_due_limpio and not v_due: st.error("❌ El RUT del Dueño ingresado no es válido.")
                else:
                    try:
                        patente_limpia = patente_v.upper().replace("-","").replace(" ","").strip()

                        con = conectar_bd()
                        cur = con.cursor()
                        
                        # Usamos r_cli directamente (que ya viene con un solo guion de la funcion validar_rut)
                        cur.execute("SELECT razon_social FROM clientes WHERE rut=%s", (r_cli,))
                        cli = cur.fetchone()
                        
                        if not cli:
                            st.error(f"❌ El cliente con RUT {r_cli} no está registrado. Ingrésalo primero en la pestaña 'Clientes'.")
                        else:
                            cur.execute("""
                                INSERT INTO vehiculos (
                                    patente, marca, modelo, año, kilometraje, nro_chasis, nro_motor, nro_puertas, color,
                                    ubicacion, observaciones, estado_reserva, dias_reserva, monto_reserva, vendedora_reserva,
                                    rut_cliente, nombre_dueno_vehiculo, rut_dueno_vehiculo, direccion_dueno_vehiculo, telefono_dueno_vehiculo, email_dueno_vehiculo,
                                    valor, comision, transferencia, valor_final
                                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                            """, (
                                patente_limpia, marca_v.upper() or None, modelo_v.upper() or None, año_v, km_v,
                                chasis_v or None, motor_v or None, puertas_v, color_v.upper() or None,
                                ubicacion_v or None, observaciones_v or None, estado_reserva_v, dias_reserva_v, monto_reserva_v, vendedora_v or None,
                                r_cli, nombre_dueno_v.upper() or None, r_due or None, dir_dueno_v or None, tel_dueno_v or None, email_dueno_v or None,
                                valor_v, comision_v, transferencia_v, valor_final_v
                            ))
                            con.commit()
                            st.success(f"✅ ¡Vehículo **{patente_limpia}** registrado a nombre de {cli[0]}!")
                    except pg_errors.UniqueViolation: 
                        st.error("❌ Error: Esa patente ya se encuentra registrada.")
                    except Exception as e: 
                        st.error(f"❌ Error: {e}")
                    finally:
                        if "cur" in locals(): cur.close()
                        if "con" in locals(): con.close()

    # --------------------------------------------------------
    # MÓDULO B: MODIFICAR VEHÍCULO EXISTENTE
    # --------------------------------------------------------
    elif st.session_state["menu_vehiculos"] == "modificar":
        st.subheader("Modificar Vehículo Existente")
        pat_key_vehiculo = f"pat_mod_vehiculo_{st.session_state.get('_gen_vehiculo', 0)}"
        pat_mod = st.text_input("Ingresa la Patente a buscar:", key=pat_key_vehiculo)
        
        if st.button("🔎 Buscar Vehículo", use_container_width=True):
            if not pat_mod.strip(): st.warning("⚠️ Ingresa una patente.")
            else:
                try:
                    con = conectar_bd()
                    cur = con.cursor()
                    cur.execute("SELECT * FROM vehiculos WHERE patente = %s", (pat_mod.strip().upper(),))
                    cols = [d[0] for d in cur.description]
                    fila = cur.fetchone()
                    if fila:
                        st.session_state["vehiculo_mod"] = dict(zip(cols, fila))
                    else: 
                        st.error("❌ Vehículo no encontrado.")
                except Exception as e: st.error(f"❌ Error: {e}")
                finally:
                    if "cur" in locals(): cur.close()
                    if "con" in locals(): con.close()

        if "vehiculo_mod" in st.session_state:
            v = st.session_state["vehiculo_mod"]
            st.success("✅ Vehículo encontrado. Modifica los campos abajo:")
            
            mv1, mv2 = st.columns(2)
            with mv1:
                marca_m = st.text_input("Marca", value=v.get("marca") or "")
                modelo_m = st.text_input("Modelo", value=v.get("modelo") or "")
                año_m = st.number_input("Año", value=int(v.get("año") or 2026), min_value=1900, max_value=2100)
                km_m = st.number_input("Kilometraje", value=int(v.get("kilometraje") or 0), min_value=0)
            with mv2:
                chasis_m = st.text_input("Nro Chasis", value=v.get("nro_chasis") or "")
                motor_m = st.text_input("Nro Motor", value=v.get("nro_motor") or "")
                puertas_m = st.number_input("Puertas", value=int(v.get("nro_puertas") or 4), min_value=0, max_value=10)
                color_m = st.text_input("Color", value=v.get("color") or "")

            st.markdown("---")
            st.markdown("**📍 Ubicación y Observaciones**")
            um1, um2 = st.columns(2)
            with um1:
                ubicacion_m = st.text_input("Ubicación", value=v.get("ubicacion") or "", key="ubicacion_mod")
            with um2:
                observaciones_m = st.text_area("Observaciones", value=v.get("observaciones") or "", height=100, key="observaciones_mod")

            st.markdown("**🔖 Estado del Vehículo**")
            estado_reserva_m = st.radio(
                "Estado", ["Disponible", "Reservado"], horizontal=True,
                index=0 if (v.get("estado_reserva") or "Disponible") == "Disponible" else 1,
                key=f"estado_reserva_mod_{v['patente']}",
            )
            if estado_reserva_m == "Reservado":
                rm1, rm2, rm3 = st.columns(3)
                with rm1:
                    dias_reserva_m = st.number_input(
                        "Días de Reserva", value=int(v.get("dias_reserva") or 0), min_value=0, step=1, key="dias_reserva_mod"
                    )
                with rm2:
                    monto_reserva_m = st.number_input(
                        "Monto de la Reserva ($)", value=int(v.get("monto_reserva") or 0), min_value=0, step=1000, key="monto_reserva_mod"
                    )
                with rm3:
                    vendedora_m = st.text_input("Vendedora que atendió", value=v.get("vendedora_reserva") or "", key="vendedora_mod")
            else:
                dias_reserva_m = None
                monto_reserva_m = None
                vendedora_m = ""

            st.markdown("---")
            st.markdown("### 💰 Modificar Valores Económicos")
            col_m_eco1, col_m_eco2, col_m_eco3 = st.columns(3)
            
            with col_m_eco1:
                val_inicial = f"{int(v.get('valor') or 0):,}".replace(",", ".")
                txt_valor_m = st.text_input(
                    "Valor Vehículo ($)", value=val_inicial, key="mod_txt_val",
                    on_change=formatear_monto_callback, args=("mod_txt_val",),
                )
                valor_m = int(txt_valor_m.replace(".", "").replace("$", "").strip() or 0)
                st.caption(f"🔎 Sombra: **${valor_m:,.0f}**".replace(",", "."))

            with col_m_eco2:
                com_inicial = f"{int(v.get('comision') or 0):,}".replace(",", ".")
                txt_comision_m = st.text_input(
                    "Comisión ($)", value=com_inicial, key="mod_txt_com",
                    on_change=formatear_monto_callback, args=("mod_txt_com",),
                )
                comision_m = int(txt_comision_m.replace(".", "").replace("$", "").strip() or 0)
                st.caption(f"🔎 Sombra: **${comision_m:,.0f}**".replace(",", "."))

            with col_m_eco3:
                trans_inicial = f"{int(v.get('transferencia') or 0):,}".replace(",", ".")
                txt_transferencia_m = st.text_input(
                    "Transferencia ($)", value=trans_inicial, key="mod_txt_trans",
                    on_change=formatear_monto_callback, args=("mod_txt_trans",),
                )
                transferencia_m = int(txt_transferencia_m.replace(".", "").replace("$", "").strip() or 0)
                st.caption(f"🔎 Sombra: **${transferencia_m:,.0f}**".replace(",", "."))
            
            valor_final_m = valor_m + comision_m + transferencia_m
            st.warning(f"💵 **Nuevo Valor Final:** ${valor_final_m:,.0f}".replace(",", "."))
            st.markdown("---")

            st.markdown("**📋 Modificar Datos del Dueño**")
            nombre_d_m = st.text_input("Nombre Dueño", value=v.get("nombre_dueno_vehiculo") or "")
            rut_d_m = st.text_input("RUT Dueño (Acepta con o sin guion)", value=v.get("rut_dueno_vehiculo") or "")
            dir_d_m = st.text_input("Dirección Dueño", value=v.get("direccion_dueno_vehiculo") or "")
            tel_d_m = st.text_input("Teléfono Dueño", value=v.get("telefono_dueno_vehiculo") or "")
            email_d_m = st.text_input("E-Mail Dueño", value=v.get("email_dueno_vehiculo") or "")

            st.write("")
            if st.button("💾 Guardar Cambios y Actualizar Vehículo", use_container_width=True, type="primary"):
                rut_d_limpio = rut_d_m.replace(".","").replace("-","").strip()
                v_due, r_due = validar_rut(rut_d_limpio) if rut_d_limpio else (True, None)
                
                if rut_d_limpio and not v_due: 
                    st.error("❌ El RUT del Dueño ingresado no es válido.")
                else:
                    try:
                        con = conectar_bd()
                        cur = con.cursor()
                        cur.execute("""
                            UPDATE vehiculos SET
                                marca=%s, modelo=%s, año=%s, kilometraje=%s, nro_chasis=%s, nro_motor=%s, nro_puertas=%s, color=%s,
                                ubicacion=%s, observaciones=%s, estado_reserva=%s, dias_reserva=%s, monto_reserva=%s, vendedora_reserva=%s,
                                nombre_dueno_vehiculo=%s, rut_dueno_vehiculo=%s, direccion_dueno_vehiculo=%s, telefono_dueno_vehiculo=%s, email_dueno_vehiculo=%s,
                                valor=%s, comision=%s, transferencia=%s, valor_final=%s
                            WHERE patente=%s
                        """, (
                            marca_m.upper() or None, modelo_m.upper() or None, año_m, km_m, chasis_m or None, motor_m or None, puertas_m, color_m.upper() or None,
                            ubicacion_m or None, observaciones_m or None, estado_reserva_m, dias_reserva_m, monto_reserva_m, vendedora_m or None,
                            nombre_d_m.upper() or None, r_due or None, dir_d_m or None, tel_d_m or None, email_d_m or None,
                            valor_m, comision_m, transferencia_m, valor_final_m, v["patente"]
                        ))
                        con.commit()
                        del st.session_state["vehiculo_mod"]
                        st.session_state["_gen_vehiculo"] = st.session_state.get("_gen_vehiculo", 0) + 1
                        st.session_state["_flash_vehiculos"] = "✅ ¡Registro actualizado con éxito en la nube!"
                        st.rerun(scope="fragment")
                    except Exception as e: st.error(f"❌ Error al actualizar: {e}")
                    finally:
                        if "cur" in locals(): cur.close()
                        if "con" in locals(): con.close()

if "vehiculos" in tabs:
    with tabs["vehiculos"]:
        render_vehiculos()

# ================================================================
# PESTAÑA 6 — BÚSQUEDA Y LISTADO
# ================================================================
@st.fragment
def render_buscar():
    st.subheader("Buscar y Listar Registros")
    tipo_busqueda = st.selectbox("¿Qué deseas buscar?", ["Clientes", "Vehículos", "Créditos", "Productos"])
    texto_buscar = st.text_input("Buscar por RUT, nombre o patente:")

    if st.button("🔎 Buscar"):
        try:
            con = conectar_bd()
            cur = con.cursor()
            if tipo_busqueda == "Clientes":
                cur.execute("""
                    SELECT rut, razon_social, giro, telefono, email, ciudad FROM clientes
                    WHERE rut ILIKE %s OR razon_social ILIKE %s ORDER BY razon_social LIMIT 50
                """, (f"%{texto_buscar}%", f"%{texto_buscar}%"))
                cols = ["RUT", "Razón Social", "Giro", "Teléfono", "E-Mail", "Ciudad"]
            elif tipo_busqueda == "Vehículos":
                cur.execute("""
                    SELECT v.patente, v.marca, v.modelo, v.año, v.kilometraje, v.color,
                           v.ubicacion, v.estado_reserva, v.dias_reserva, v.monto_reserva, v.vendedora_reserva, v.observaciones,
                           v.rut_cliente, c.razon_social
                    FROM vehiculos v LEFT JOIN clientes c ON c.rut = v.rut_cliente
                    WHERE v.patente ILIKE %s OR v.marca ILIKE %s OR v.rut_cliente ILIKE %s ORDER BY v.patente LIMIT 50
                """, (f"%{texto_buscar}%", f"%{texto_buscar}%", f"%{texto_buscar}%"))
                cols = [
                    "Patente", "Marca", "Modelo", "Año", "KM", "Color",
                    "Ubicación", "Estado", "Días Reserva", "Monto Reserva", "Vendedora", "Observaciones",
                    "RUT Cliente", "Cliente",
                ]
            elif tipo_busqueda == "Créditos":
                cur.execute("""
                    SELECT c.patente, c.nombre_cliente, c.rut_cliente, c.precio_venta, c.pie, c.saldo_financiamiento,
                           c.cantidad_cuotas, c.cuotas_pendientes, c.monto_comision, c.comision_pagada, c.forma_pago,
                           (SELECT COUNT(*) FROM documentos_credito d WHERE d.credito_id = c.id)
                    FROM creditos c
                    WHERE c.patente ILIKE %s OR c.rut_cliente ILIKE %s OR c.nombre_cliente ILIKE %s
                    ORDER BY c.creado_en DESC LIMIT 50
                """, (f"%{texto_buscar}%", f"%{texto_buscar}%", f"%{texto_buscar}%"))
                cols = [
                    "Patente", "Cliente", "RUT Cliente", "Precio Venta", "Pie", "Saldo Financiamiento",
                    "N° Cuotas", "Cuotas Pendientes", "Monto Comisión", "Comisión Pagada", "Forma de Pago",
                    "N° Documentos",
                ]
            else:
                cur.execute("""
                    SELECT codigo_producto, codigo_barra, descripcion, stock_actual, stock_minimo,
                           stock_reposicion, stock_maximo, costo, margen_venta
                    FROM productos
                    WHERE codigo_producto ILIKE %s OR codigo_barra ILIKE %s OR descripcion ILIKE %s
                    ORDER BY descripcion LIMIT 50
                """, (f"%{texto_buscar}%", f"%{texto_buscar}%", f"%{texto_buscar}%"))
                cols = [
                    "Código Producto", "Código Barra", "Descripción", "Stock Actual", "Stock Mínimo",
                    "Stock Reposición", "Stock Máximo", "Costo", "Margen %",
                ]

            filas = cur.fetchall()
            if filas:
                st.dataframe(pd.DataFrame(filas, columns=cols), use_container_width=True)
                st.caption(f"{len(filas)} resultado(s) encontrado(s).")
            else:
                st.info("No se encontraron resultados.")
        except Exception as e: st.error(f"❌ Error: {e}")
        finally:
            if "cur" in locals(): cur.close()
            if "con" in locals(): con.close()

if "buscar" in tabs:
    with tabs["buscar"]:
        render_buscar()

# ================================================================
# PESTAÑA — ADMINISTRACIÓN (solo visible para el Administrador)
# ================================================================
@st.fragment
def render_admin():
    if st.session_state.get("_flash_admin"):
        st.success(st.session_state.pop("_flash_admin"))

    st.subheader("⚙️ Administración de Usuarios")
    st.caption("Esta pestaña solo la ven los usuarios con perfil Administrador.")

    try:
        con = conectar_bd()
        cur = con.cursor()
        cur.execute(
            "SELECT id, username, nombre_completo, perfil, es_admin, pestañas_permitidas, activo "
            "FROM usuarios ORDER BY username"
        )
        cols = [d[0] for d in cur.description]
        usuarios_lista = [dict(zip(cols, fila)) for fila in cur.fetchall()]
    except Exception as e:
        st.error(f"❌ Error al cargar usuarios: {e}")
        usuarios_lista = []
    finally:
        if "cur" in locals(): cur.close()
        if "con" in locals(): con.close()

    st.markdown("### 👥 Usuarios existentes")
    if usuarios_lista:
        tabla = [{
            "Usuario": u["username"],
            "Nombre": u["nombre_completo"] or "",
            "Perfil": u["perfil"],
            "Administrador": "Sí" if u["es_admin"] else "No",
            "Pestañas": "Todas" if u["es_admin"] else (u["pestañas_permitidas"] or "—"),
            "Activo": "Sí" if u["activo"] else "No",
        } for u in usuarios_lista]
        st.dataframe(pd.DataFrame(tabla), use_container_width=True)
    else:
        st.info("No hay usuarios registrados todavía.")

    st.markdown("---")
    st.markdown("### ➕ Crear nuevo usuario")

    nuevo_username = st.text_input("Nombre de usuario *", key="admin_nuevo_username")
    nc1, nc2 = st.columns(2)
    with nc1:
        nuevo_nombre = st.text_input("Nombre completo", key="admin_nuevo_nombre")
        nuevo_perfil = st.text_input("Perfil (ej: Taller, Vendedor)", key="admin_nuevo_perfil")
    with nc2:
        nueva_clave = st.text_input("Contraseña *", type="password", key="admin_nueva_clave")
        nuevo_es_admin = st.checkbox("Es Administrador (acceso a todo, incluida esta pestaña)", key="admin_nuevo_es_admin")

    st.markdown("**Pestañas a las que tendrá acceso** _(no aplica si es Administrador, ya que ve todas)_")
    cols_check = st.columns(len(TABS_DISPONIBLES))
    seleccion_pestañas = []
    for i, (clave_tab, etiqueta_tab) in enumerate(TABS_DISPONIBLES):
        with cols_check[i]:
            if st.checkbox(etiqueta_tab, key=f"admin_check_{clave_tab}"):
                seleccion_pestañas.append(clave_tab)

    if st.button("💾 Crear Usuario", type="primary"):
        if not nuevo_username.strip() or not nueva_clave.strip():
            st.warning("⚠️ El nombre de usuario y la contraseña son obligatorios.")
        elif not nuevo_es_admin and not seleccion_pestañas:
            st.warning("⚠️ Selecciona al menos una pestaña, o marca a este usuario como Administrador.")
        else:
            try:
                con = conectar_bd()
                cur = con.cursor()
                hash_nuevo, salt_nuevo = hash_password(nueva_clave)
                cur.execute("""
                    INSERT INTO usuarios (username, password_hash, password_salt, nombre_completo, perfil, es_admin, pestañas_permitidas)
                    VALUES (%s,%s,%s,%s,%s,%s,%s)
                """, (
                    nuevo_username.strip(), hash_nuevo, salt_nuevo, nuevo_nombre.strip() or None,
                    nuevo_perfil.strip() or "Usuario", nuevo_es_admin, ",".join(seleccion_pestañas),
                ))
                con.commit()
                st.session_state["_flash_admin"] = f"✅ Usuario '{nuevo_username.strip()}' creado correctamente."
                st.rerun(scope="fragment")
            except pg_errors.UniqueViolation:
                st.error("❌ Ese nombre de usuario ya existe.")
            except Exception as e:
                st.error(f"❌ Error: {e}")
            finally:
                if "cur" in locals(): cur.close()
                if "con" in locals(): con.close()

    st.markdown("---")
    st.markdown("### ✏️ Modificar usuario existente")

    if usuarios_lista:
        opciones_usuario = {f"{u['username']} ({u['perfil']})": u for u in usuarios_lista}
        seleccion_label = st.selectbox("Selecciona un usuario", list(opciones_usuario.keys()), key="admin_editar_select")
        u_sel = opciones_usuario[seleccion_label]

        ec1, ec2 = st.columns(2)
        with ec1:
            editar_nombre = st.text_input("Nombre completo", value=u_sel.get("nombre_completo") or "", key="admin_editar_nombre")
            editar_perfil = st.text_input("Perfil", value=u_sel.get("perfil") or "", key="admin_editar_perfil")
            editar_activo = st.checkbox("Usuario activo", value=u_sel.get("activo", True), key="admin_editar_activo")
        with ec2:
            editar_es_admin = st.checkbox(
                "Es Administrador", value=u_sel.get("es_admin", False), key="admin_editar_es_admin"
            )
            nueva_clave_reset = st.text_input(
                "Nueva contraseña (deja vacío para no cambiarla)", type="password", key="admin_editar_clave"
            )

        st.markdown("**Pestañas a las que tiene acceso**")
        permitidas_actuales = set((u_sel.get("pestañas_permitidas") or "").split(","))
        cols_edit = st.columns(len(TABS_DISPONIBLES))
        seleccion_editar = []
        for i, (clave_tab, etiqueta_tab) in enumerate(TABS_DISPONIBLES):
            with cols_edit[i]:
                marcado = clave_tab in permitidas_actuales
                if st.checkbox(etiqueta_tab, value=marcado, key=f"admin_editar_check_{clave_tab}_{u_sel['id']}"):
                    seleccion_editar.append(clave_tab)

        eb1, eb2 = st.columns(2)
        with eb1:
            guardar_usuario = st.button("💾 Guardar Cambios", type="primary", use_container_width=True)
        with eb2:
            eliminar_usuario = st.button("🗑️ Eliminar Usuario", use_container_width=True)

        if guardar_usuario:
            try:
                con = conectar_bd()
                cur = con.cursor()
                if nueva_clave_reset.strip():
                    hash_edit, salt_edit = hash_password(nueva_clave_reset)
                    cur.execute("""
                        UPDATE usuarios SET nombre_completo=%s, perfil=%s, es_admin=%s, pestañas_permitidas=%s,
                            activo=%s, password_hash=%s, password_salt=%s
                        WHERE id=%s
                    """, (
                        editar_nombre.strip() or None, editar_perfil.strip() or "Usuario", editar_es_admin,
                        ",".join(seleccion_editar), editar_activo, hash_edit, salt_edit, u_sel["id"],
                    ))
                else:
                    cur.execute("""
                        UPDATE usuarios SET nombre_completo=%s, perfil=%s, es_admin=%s, pestañas_permitidas=%s, activo=%s
                        WHERE id=%s
                    """, (
                        editar_nombre.strip() or None, editar_perfil.strip() or "Usuario", editar_es_admin,
                        ",".join(seleccion_editar), editar_activo, u_sel["id"],
                    ))
                con.commit()
                st.session_state["_flash_admin"] = f"✅ Usuario '{u_sel['username']}' actualizado."
                st.rerun(scope="fragment")
            except Exception as e:
                st.error(f"❌ Error: {e}")
            finally:
                if "cur" in locals(): cur.close()
                if "con" in locals(): con.close()

        if eliminar_usuario:
            if u_sel["username"] == usuario_actual["username"]:
                st.error("❌ No puedes eliminar el usuario con el que iniciaste sesión.")
            else:
                try:
                    con = conectar_bd()
                    cur = con.cursor()
                    cur.execute("DELETE FROM usuarios WHERE id=%s", (u_sel["id"],))
                    con.commit()
                    st.session_state["_flash_admin"] = f"✅ Usuario '{u_sel['username']}' eliminado."
                    st.rerun(scope="fragment")
                except Exception as e:
                    st.error(f"❌ Error: {e}")
                finally:
                    if "cur" in locals(): cur.close()
                    if "con" in locals(): con.close()
    else:
        st.info("Aún no hay usuarios para modificar.")

if "admin" in tabs:
    with tabs["admin"]:
        render_admin()