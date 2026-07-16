"""Generación de documentos Word a partir de plantillas con marcadores [CLAVE]."""
import os

from docx import Document

PLANTILLAS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Plantillas")

PLANTILLA_COMPRAVENTA = os.path.join(PLANTILLAS_DIR, "Plantilla_Compraventa.docx")


def _unir_runs_con_marcador(parrafo, marcadores):
    """Word puede partir un marcador como '[NOM' + 'BRE]' entre varios runs.
    Si el texto completo del párrafo contiene el marcador pero ningún run individual
    lo tiene completo, fusionamos todos los runs en uno para poder reemplazar.
    """
    texto_completo = parrafo.text
    necesita_fusion = any(
        m in texto_completo and not any(m in run.text for run in parrafo.runs)
        for m in marcadores
    )
    if necesita_fusion and parrafo.runs:
        primer_run = parrafo.runs[0]
        primer_run.text = texto_completo
        for run in parrafo.runs[1:]:
            run.text = ""


def reemplazar_en_parrafo(parrafo, mapa):
    _unir_runs_con_marcador(parrafo, mapa.keys())
    for clave, valor in mapa.items():
        if clave in parrafo.text:
            for run in parrafo.runs:
                if clave in run.text:
                    run.text = run.text.replace(clave, str(valor) if valor else "")


def reemplazar_en_documento(doc, mapa):
    for p in doc.paragraphs:
        reemplazar_en_parrafo(p, mapa)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    reemplazar_en_parrafo(p, mapa)


def generar_docx_desde_plantilla(ruta_plantilla, reemplazos):
    """Devuelve un BytesIO con el .docx generado a partir de la plantilla y los reemplazos."""
    import io

    doc = Document(ruta_plantilla)
    reemplazar_en_documento(doc, reemplazos)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
